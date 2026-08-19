"""

Módulo 3: Generador de DJC (Declaración Jurada de Conformidad)

==============================================================

Orquestador principal. La lógica pesada está modularizada en:

  - modules/extractors/   → Extractores por OEC

  - modules/regulations.py → Detección de reglamentos

  - modules/pdf_ops.py    → Censura y merge de PDFs

"""

from __future__ import annotations



import os

import re

import json

import logging

import zoneinfo

from datetime import datetime, timedelta

from pathlib import Path

from typing import Optional



import fitz  # PyMuPDF  # type: ignore[import-untyped]

from docx import Document  # type: ignore[import-untyped]



from modules.regulations import detect_reglamento

from modules.extractors.dispatcher import detect_oec, extract_product_data

from modules.pdf_ops import censor_cert_pdf, merge_pdfs, strip_old_djc, extract_pdf_clean_text


from modules.extractors.shared import parse_date, calc_vencimiento, calc_inicio_tramite



logger = logging.getLogger(__name__)



def normalize_oec_key(oec_key: str) -> str:
    """Normaliza las variaciones del nombre OEC (como Quektra/Qetkra) a la clave estándar de m3_config.json."""
    if not oec_key:
        return ""
    k_lower = oec_key.lower().strip()
    if k_lower in ("quektra", "qetkra"):
        return "Qetkra"
    return oec_key


class DJCGenerator:

    """Genera Declaraciones Juradas de Conformidad a partir de datos de certificados."""



    TEMPLATE_FILENAME = "DJ Conformidad Modelo SE.docx"



    def __init__(self, config_path: Optional[str] = None, gui_logger=None):

        self.gui_logger = gui_logger



        if config_path is None:

            base_dir = Path(__file__).parent.parent

            config_path = str(base_dir / "m3_config.json")



        with open(config_path, "r", encoding="utf-8") as f:

            self.config = json.load(f)



        base_dir = Path(__file__).parent.parent
        self.template_path = base_dir / "assets" / "djc_templates" / self.TEMPLATE_FILENAME
        if not self.template_path.exists():
            self.template_path = base_dir / self.TEMPLATE_FILENAME
            if not self.template_path.exists():
                raise FileNotFoundError(f"Template DJC no encontrado en assets/djc_templates ni en raíz: {self.TEMPLATE_FILENAME}")

        self._log("info", f"DJCGenerator inicializado. Template: {self.template_path}")



    # ─── Logging dual ──────────────────────────────────────



    def _log(self, level: str, message: str):

        getattr(logger, level, logger.info)(message)

        if self.gui_logger:

            try:

                self.gui_logger.log(message, level.upper())

            except Exception:

                pass



    def _log_fn(self, level: str, msg: str):

        """Adaptador para pasar a funciones externas."""

        self._log(level, msg)



    # ─── Delegación a módulos especializados ──────────────



    def detect_reglamento(self, normas_text: str, producto_desc: str = "") -> str:

        return detect_reglamento(normas_text, producto_desc, self._log_fn)



    def detect_oec(self, cert_text: str) -> str:

        return detect_oec(cert_text, self._log_fn)



    def censor_cert_pdf(self, doc, fabricante: str = "", direccion: str = "",

                        preserve_words: Optional[list[str]] = None):

        return censor_cert_pdf(doc, fabricante, direccion, preserve_words, self._log_fn)



    def merge_pdfs(self, djc_pdf_path: str, cert_pdf_path: str,

                   output_path: Optional[str] = None,

                   extra_pdfs: Optional[list[str]] = None) -> str:

        return merge_pdfs(djc_pdf_path, cert_pdf_path, output_path, extra_pdfs, self._log_fn)



    def _strip_old_djc(self, doc):

        strip_old_djc(doc, self._log_fn)



    # ─── Extracción desde PDF ──────────────────────────────



    def extract_cert_data(self, pdf_path: str) -> dict:

        """Extrae datos clave de un certificado PDF (número, normas, OEC, fechas)."""

        if not os.path.exists(pdf_path):

            raise FileNotFoundError(f"Certificado no encontrado: {pdf_path}")



        full_text_sorted = extract_pdf_clean_text(pdf_path)
        doc = fitz.open(pdf_path)
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        doc.close()

        if not full_text_sorted.strip():
            full_text_sorted = full_text




        # Detectar PDF sin texto extraíble (posible escaneado sin OCR)

        if not full_text.strip():

            self._log("warning", "[M3] ⚠️ PDF sin texto extraíble (posible escaneado sin OCR). "

                      "La IA recibirá texto vacío y no podrá extraer campos.")

        else:

            char_count = len(full_text.strip())

            self._log("info", f"[M3] Texto extraído del PDF: {char_count:,} caracteres")





        result = {

            "cert_number":     self._extract_cert_number(full_text),

            "normas":          self._extract_normas(full_text),

            "fecha_emision":   self._extract_date(full_text, "emision"),

            "fecha_vencimiento": self._extract_date(full_text, "vencimiento"),

            "oec_key":         self.detect_oec(full_text),

            "cert_text":       full_text,

            "cert_text_sorted": full_text_sorted,

        }

        self._log("info", f"Datos extraídos del certificado: Nro={result['cert_number']}")

        return result



    def extract_product_data_from_cert(self, text: str, text_sorted: str = "") -> dict:

        """Dispatcher de extracción de datos del producto (delega a extractors/)."""

        oec_key = self.detect_oec(text)

        return extract_product_data(text, text_sorted, oec_key, self._log_fn)



    # ─── Preparación de datos ─────────────────────────────



    def prepare_from_certificate(self, cert_pdf_path: str) -> dict:

        """

        Prepara datos DJC desde un certificado nuevo (sin auditoría previa).

        Extrae TODOS los datos del producto directamente del PDF.

        """

        cert_data = self.extract_cert_data(cert_pdf_path)

        text = cert_data.get("cert_text", "")

        text_sorted = cert_data.get("cert_text_sorted", "")



        normas = cert_data.get("normas", "")

        reglamento = self.detect_reglamento(normas)

        oec_key = cert_data.get("oec_key", "")
        normalized_oec_key = normalize_oec_key(oec_key)
        oec_info = self.config["oec_options"].get(normalized_oec_key, {})

        cert_number = cert_data.get("cert_number", "")



        self._log("info", f"[M3] OEC={oec_key or '[no detectado]'}, Cert={cert_number or '[no encontrado]'}, Reglamento={reglamento or '[no detectado]'}")



        product = self.extract_product_data_from_cert(text, text_sorted)



        fecha_emision = product.get("fecha_emision", "") or cert_data.get("fecha_emision", "")

        fecha_vencimiento = product.get("fecha_vencimiento", "") or cert_data.get("fecha_vencimiento", "")

        fecha_inicio_tramite = product.get("fecha_inicio_tramite", "")



        if not fecha_vencimiento and fecha_emision:

            fecha_vencimiento = calc_vencimiento(fecha_emision, reglamento)

        if not fecha_inicio_tramite and fecha_vencimiento:

            fecha_inicio_tramite = calc_inicio_tramite(fecha_vencimiento)



        data = {

            "djc_id":                   self.generate_djc_id(reglamento, oec_key, bidcom_num=""),

            "oec_key":                  oec_key,

            "fabricante":               product["fabricante"],

            "direccion_fabrica":        product["direccion"],

            "producto_desc":            product["producto_desc"],

            "marca":                    product["marca"],

            "modelos":                  product["modelos"],

            "specs":                    product["specs"],

            "reglamento":               reglamento,

            "normas":                   normas,

            "cert_number":              cert_number,

            "esquema":                  self.config["esquema_options"][0],

            "fecha_emision":            fecha_emision,

            "fecha_vigilancia":         "No aplicable",

            "fecha_proxima_vigilancia": fecha_vencimiento,

            "fecha_inicio_tramite":     fecha_inicio_tramite,

            "oec_nombre":               oec_info.get("nombre", oec_key),

            "oec_contacto":             oec_info.get("contacto", ""),

            "enlace_djc":               "https://qr.gadnic.com/certifications/certificado-",

        }



        self._log("info", f"Datos extraídos: marca={product['marca']}, emision={fecha_emision}, vencimiento={fecha_vencimiento}")

        self._log("info", f"[M3] DJC ID propuesto: {data['djc_id']}")

        return data



    def prepare_from_audit(self, json_data: dict, cert_data: dict) -> dict:

        """Prepara datos DJC desde resultado de auditoría (Módulo 2 → OK)."""

        normas = cert_data.get("normas", "")

        producto_desc = json_data.get("specs_tecnicas", "")

        if isinstance(producto_desc, list):

            producto_desc = ", ".join(producto_desc)



        reglamento = self.detect_reglamento(normas, producto_desc)

        oec_key = cert_data.get("oec_key", "")
        normalized_oec_key = normalize_oec_key(oec_key)
        oec_info = self.config["oec_options"].get(normalized_oec_key, {})



        fecha_emision = cert_data.get("fecha_emision", "")

        fecha_vencimiento = cert_data.get("fecha_vencimiento", "")

        if not fecha_vencimiento and fecha_emision:

            fecha_vencimiento = calc_vencimiento(fecha_emision, reglamento)

        fecha_inicio_tramite = calc_inicio_tramite(fecha_vencimiento)



        bidcom_number = json_data.get("id_gestion", "")

        cert_ref = cert_data.get("cert_number", "")

        djc_id = f"DJC-CERTIFICADO {bidcom_number}-V1" if bidcom_number else "DJC-V1"



        modelos = json_data.get("modelos_solicitados", [])

        modelos_str = ", ".join(modelos) if isinstance(modelos, list) else str(modelos)



        specs = json_data.get("specs_tecnicas", "")

        if isinstance(specs, list):

            specs = ", ".join(specs)



        return {

            "djc_id":                   djc_id,

            "bidcom_number":            bidcom_number,

            "cert_number":              cert_ref,

            "fabricante":               json_data.get("fabrica", ""),

            "direccion_fabrica":        json_data.get("direccion_fabrica", ""),

            "producto_desc":            self._infer_product_description(json_data),

            "marca":                    json_data.get("marca", ""),

            "modelos":                  modelos_str,

            "specs":                    specs,

            "reglamento":               reglamento,

            "normas":                   normas,

            "esquema":                  self.config["esquema_options"][0],

            "fecha_emision":            fecha_emision,

            "fecha_vigilancia":         "No aplicable",

            "fecha_proxima_vigilancia": fecha_vencimiento,

            "fecha_inicio_tramite":     fecha_inicio_tramite,

            "oec_nombre":               oec_info.get("nombre", oec_key),

            "oec_contacto":             oec_info.get("contacto", ""),

            "enlace_djc":               self._generate_djc_link(bidcom_number),

        }



    # ─── Generación de ID ──────────────────────────────────



    def generate_djc_id(self, reglamento: str, oec_nombre: str, bidcom_num: Optional[str] = None) -> str:

        """Genera el código ID de la DJC: DJC-{REG}-{MMYY}-{BIDCOM}-{OEC}-V1."""

        regl_abrev_map = [

            (["juguete", "163/2004", "nm 300"],                                   "SJ"),

            (["16/2025", "17/2025", "60335", "62368", "62841", "62040", "60065"], "SE"),

            (["eficiencia energ", "mínima eficien"],                               "EE"),

            (["biciclet", "nm 301"],                                               "BI"),

            (["anteojos", "iso 12312"],                                            "AO"),

            (["encendedor", "iso 9994", "iram 3980"],                              "EN"),

            (["ftalato", "583/2008"],                                              "FT"),

        ]

        reglamento_raw = (reglamento or "").lower()

        regl_abrev = "OT"

        for keywords, code in regl_abrev_map:

            if any(kw in reglamento_raw for kw in keywords):

                regl_abrev = code

                break



        oec_abrev_map = {
            "Lenor":          "LNR",
            "Quektra":        "QKA",
            "Qetkra":         "QKA",
            "Intertek":       "ITK",
            "Bureau Veritas": "BVA",
            "TÜV":            "TUV",
            "IRAM":           "IRM",
        }

        oec_raw = (oec_nombre or "").strip()

        oec_abrev = next(

            (v for k, v in oec_abrev_map.items() if k.lower() in oec_raw.lower()),

            "ORG"

        )



        anio_mes = datetime.now().strftime("%m%y")

        # Formato con prefijo C igual que el frontend: C888, C912, etc.

        bidcom_clean = str(bidcom_num).replace("/", "-").replace("C", "").replace("c", "") if bidcom_num else ""

        bidcom_str = f"C{bidcom_clean}" if bidcom_clean else "XXXX"

        return f"DJC-{regl_abrev}-{anio_mes}-{bidcom_str}-{oec_abrev}-V1"



    # ─── Llenado de plantilla Word ─────────────────────────



    def fill_template(self, data: dict) -> Document:

        """Llena la plantilla DJC Word con los datos proporcionados."""

        doc = Document(str(self.template_path))

        tables = doc.tables



        if len(tables) < 8:

            raise ValueError(f"Template inválido: esperaba 8 tablas, encontró {len(tables)}")



        emp = data.get("empresa_override") or self.config["empresa"]

        self._set_cell_id(tables[0], 1, 0, data.get("djc_id", ""))



        self._set_cell(tables[1], 0, 1, emp.get("razon_social", ""))

        self._set_cell(tables[1], 1, 1, emp.get("cuit", ""))

        self._set_cell(tables[1], 2, 1, emp.get("marca_registrada", ""))

        self._set_cell(tables[1], 3, 1, emp.get("domicilio_legal", ""))

        self._set_cell(tables[1], 4, 1, emp.get("domicilio_deposito", ""))

        self._set_cell(tables[1], 5, 1, emp.get("telefono", ""))

        self._set_cell(tables[1], 6, 1, emp.get("email", ""))



        rep = data.get("representante", self.config.get("representante_autorizado")) or {}

        self._set_cell(tables[2], 0, 1, rep.get("nombre", "No Aplica"))

        self._set_cell(tables[2], 1, 1, rep.get("cuit", "No Aplica"))

        self._set_cell(tables[2], 2, 1, rep.get("domicilio", "No Aplica"))



        self._set_cell(tables[3], 0, 1, "ver «Modelo» más abajo")

        self._set_cell(tables[3], 1, 1, data.get("fabricante", ""))

        self._set_cell(tables[3], 2, 1, data.get("direccion_fabrica", ""))

        self._set_cell(tables[3], 3, 1, data.get("producto_desc", ""))

        self._set_cell(tables[3], 4, 1, data.get("marca", ""))

        self._set_cell(tables[3], 5, 1, data.get("modelos", ""))

        self._set_cell(tables[3], 6, 1, data.get("specs", ""))



        self._set_cell(tables[4], 0, 1, data.get("reglamento", ""))

        self._set_cell(tables[4], 0, 2, data.get("reglamento", ""))

        self._set_cell(tables[4], 1, 1, data.get("normas", ""))

        self._set_cell(tables[4], 1, 2, data.get("normas", ""))

        self._set_cell(tables[4], 2, 2, data.get("cert_number", ""))

        self._set_cell(tables[4], 3, 2, data.get("esquema", ""))

        self._set_cell(tables[4], 4, 2, data.get("fecha_emision", ""))

        self._set_cell(tables[4], 5, 2, data.get("fecha_vigilancia", "No aplicable"))

        self._set_cell(tables[4], 6, 2, data.get("fecha_proxima_vigilancia", ""))

        self._set_cell(tables[4], 7, 2, data.get("oec_nombre", ""))

        self._set_cell(tables[4], 8, 2, data.get("oec_nombre", ""))

        self._set_cell(tables[4], 9, 2, data.get("oec_contacto", ""))



        enlace = data.get("enlace_djc", "")

        if enlace:

            self._set_cell_hyperlink(tables[5], 0, 1, enlace, enlace)

        else:

            self._set_cell(tables[5], 0, 1, "")



        tz = zoneinfo.ZoneInfo("America/Argentina/Buenos_Aires")

        fecha_hoy = datetime.now(tz).strftime("%d/%m/%Y")

        self._set_cell(tables[6], 0, 1, data.get("fecha_emision_djc", fecha_hoy))

        self._set_cell(tables[6], 1, 1, self.config["emision"]["lugar"])

        self._set_cell(tables[7], 1, 1, self.config["firma"]["aclaracion"])



        self._log("info", f"Template DJC llenado para cert: {data.get('cert_number', 'N/A')}")

        return doc



    # ─── Helpers de celdas Word ────────────────────────────



    def _set_cell(self, table, row: int, col: int, value: str):

        from docx.shared import Pt  # type: ignore[import-untyped]

        try:

            cell = table.rows[row].cells[col]

            if cell.paragraphs:

                para = cell.paragraphs[0]

                if para.runs:

                    run = para.runs[0]

                    run.text = str(value) if value else ""

                    run.font.name = "Arial"

                    run.font.size = Pt(8)

                    for extra_run in para.runs[1:]:

                        extra_run.text = ""

                else:

                    para.clear()

                    run = para.add_run(str(value) if value else "")

                    run.font.name = "Arial"

                    run.font.size = Pt(8)

            else:

                cell.text = str(value) if value else ""

        except (IndexError, AttributeError) as e:

            self._log("warning", f"No se pudo escribir en celda [{row},{col}]: {e}")



    def _set_cell_id(self, table, row: int, col: int, value: str):

        from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

        from docx.oxml.ns import qn as _qn  # type: ignore[import-untyped]

        from lxml import etree  # type: ignore[import-untyped]

        try:

            cell = table.rows[row].cells[col]

            tc = cell._tc

            tcPr = tc.get_or_add_tcPr()

            for shd in tcPr.findall(_qn('w:shd')):

                tcPr.remove(shd)

            shd = etree.SubElement(tcPr, _qn('w:shd'))

            shd.set(_qn('w:val'), 'clear')

            shd.set(_qn('w:color'), 'auto')

            shd.set(_qn('w:fill'), 'FFFFFF')



            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

            para.clear()

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run(str(value) if value else "")

            run.font.name = 'Arial'

            run.font.size = Pt(10)

            run.font.bold = True

            run.font.color.rgb = RGBColor(0, 0, 0)

        except Exception as e:

            self._log('warning', f'[M3] No se pudo formatear celda ID [{row},{col}]: {e}')

            self._set_cell(table, row, col, value)



    def _set_cell_hyperlink(self, table, row: int, col: int, url: str, display_text: Optional[str] = None):

        try:

            from docx.oxml.shared import OxmlElement, qn  # type: ignore[import-untyped]

            cell = table.rows[row].cells[col]

            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

            para.clear()



            hyperlink = OxmlElement('w:hyperlink')

            hyperlink.set(qn('r:id'), para.part.relate_to(

                url,

                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',

                is_external=True

            ))



            new_run = OxmlElement('w:r')

            rPr = OxmlElement('w:rPr')

            color = OxmlElement('w:color')

            color.set(qn('w:val'), '0000FF')

            rPr.append(color)

            u = OxmlElement('w:u')

            u.set(qn('w:val'), 'single')

            rPr.append(u)

            rFonts = OxmlElement('w:rFonts')

            rFonts.set(qn('w:ascii'), 'Arial')

            rFonts.set(qn('w:hAnsi'), 'Arial')

            rPr.append(rFonts)

            sz = OxmlElement('w:sz')

            sz.set(qn('w:val'), '16')

            rPr.append(sz)

            szCs = OxmlElement('w:szCs')

            szCs.set(qn('w:val'), '16')

            rPr.append(szCs)



            new_run.append(rPr)

            new_run.text = display_text if display_text else url

            hyperlink.append(new_run)

            para._element.append(hyperlink)

        except (IndexError, AttributeError) as e:

            self._log("warning", f"No se pudo crear hipervínculo en celda [{row},{col}]: {e}")

            self._set_cell(table, row, col, url)



    # ─── Export ────────────────────────────────────────────



    def save_docx(self, doc: Document, output_path: str) -> str:

        doc.save(output_path)

        self._log("info", f"DJC Word guardada: {output_path}")

        return output_path



    def export_to_pdf(self, docx_path: str, pdf_path: Optional[str] = None) -> str:
        if pdf_path is None:
            pdf_path = docx_path.replace(".docx", ".pdf")

        docx_abs = os.path.abspath(docx_path)
        pdf_abs  = os.path.abspath(pdf_path)

        # ── Intento 1: MS Word via COM (control manual del ciclo de vida) ──────
        try:
            import comtypes.client  # type: ignore[import-untyped]
            import comtypes
            wdFormatPDF = 17
            word = None
            doc  = None
            try:
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(docx_abs)
                doc.SaveAs(pdf_abs, FileFormat=wdFormatPDF)
                self._log("info", f"DJC PDF generado (via MS Word): {pdf_abs}")
                return pdf_abs
            finally:
                if doc is not None:
                    try:
                        doc.Close(False)
                    except Exception:
                        pass
                if word is not None:
                    try:
                        word.Quit()
                    except Exception:
                        pass
                    try:
                        del word
                    except Exception:
                        pass
        except Exception as e:
            self._log("warning", f"docx2pdf (MS Word) falló o no está disponible: {e}. Intentando LibreOffice...")

        # ── Intento 2: LibreOffice ────────────────────────────────────────────
        import subprocess

        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",  # Si está en el PATH
        ]

        soffice_exe = None
        for path in libreoffice_paths:
            if path == "soffice":
                try:
                    subprocess.run([path, "--version"], capture_output=True, check=True)
                    soffice_exe = path
                    break
                except Exception:
                    continue
            elif os.path.exists(path):
                soffice_exe = path
                break

        if soffice_exe:
            try:
                outdir = os.path.dirname(pdf_abs)
                subprocess.run(
                    [soffice_exe, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_abs],
                    check=True,
                    capture_output=True,
                )
                lo_output_path = os.path.join(outdir, os.path.splitext(os.path.basename(docx_abs))[0] + ".pdf")
                if lo_output_path != pdf_abs and os.path.exists(lo_output_path):
                    if os.path.exists(pdf_abs):
                        os.remove(pdf_abs)
                    os.rename(lo_output_path, pdf_abs)

                self._log("info", f"DJC PDF generado (via LibreOffice): {pdf_abs}")
                return pdf_abs
            except Exception as lo_err:
                self._log("error", f"Error convirtiendo con LibreOffice: {lo_err}")
                raise RuntimeError(f"Error MS Word: {e} | Error LibreOffice: {lo_err}")  # type: ignore[reportPossiblyUnbound]
        else:
            self._log("error", "No se encontró Microsoft Word ni LibreOffice instalados en el sistema.")
            raise RuntimeError(f"No se pudo generar PDF. Error MS Word: {e}. LibreOffice no instalado.")  # type: ignore[reportPossiblyUnbound]



    # ─── Utilidades ────────────────────────────────────────



    def _extract_cert_number(self, text: str) -> str:

        patterns = [

            # Etiqueta "Ref. Cert. No." — usada por Intertek IACSA (ej: TCSE-IACSA-0146/324.1)

            r"(?:Ref\.?\s*Cert\.?\s*No\.?|Ref(?:erencia)?\s*Certificado)\s*[:\-]?\s*([A-Z]{2,6}-[A-Z]{2,6}-\d{4}/\d{1,4}(?:\.\d+)?(?:R\d+)?)",

            # Formato XXXX-YYYY-NNNN/NNN.N (TCSE-IACSA-0146/324.1) — busca en texto libre

            r"\b([A-Z]{2,6}-[A-Z]{2,6}-\d{4}/\d{1,4}(?:\.\d+)?(?:R\d+)?)\b",

            # Formato Quektra: Q-AR-XXXXX-T-0

            r"\b(Q-AR-\d{4,8}(?:-[A-Z0-9]+)*)\b",

            # CB Certificate No.

            r"(?:CB\s*Certificate\s*(?:No|N[°º])\.?\s*:?\s*)([A-Z0-9][\w\-/.]+)",

            # Certificate No / Number

            r"(?:Certificate\s*(?:No|Number|#|N[°º])\.?\s*:?\s*)([A-Z0-9][\w\-/.]+)",

            # Referencia de Certificado (Lenor)

            r"(?:Referencia\s*de\s*Certificado|Certificate\s*reference)\s*:?\s*([A-Z]{2,6}-\d{3,6})",

            # Certificado No / Nro

            r"(?:Certificado\s*(?:No|N[°º]|Nro)\.?\s*:?\s*)([A-Z0-9][\w\-/.]+)",

            # Report No

            r"(?:Report\s*(?:No|Number)\.?\s*:?\s*)([A-Z0-9][\w\-/.]+)",

            # DC-X-XXXXXX formato

            r"\b(DC-[A-Z]-[A-Z0-9]{2,6}\s*-\d+(?:\.\d+)?)\b",

            r"\b(DC-[A-Z0-9\-]{5,15})\b",

            # Genérico fallback — va último para evitar falsos positivos

            r"\b([A-Z]{2,5}[\-/]\d{3,6}(?:[\-/]\w+)?)\b",

        ]

        for pattern in patterns:

            match = re.search(pattern, text, re.IGNORECASE)

            if match:

                return match.group(1).strip()

        return ""



    def _extract_normas(self, text: str) -> str:

        patterns = [

            re.compile(r'IEC\s*(?!17\d{3})\d{4,5}(?:[\-\.\w]*)?(?:[^\n]*)', re.IGNORECASE),

            re.compile(r'CISPR\s*\d{1,3}(?:[\-\.\w]*)?(?:[^\n]*)', re.IGNORECASE),

            re.compile(r'ISO(?:/IEC)?\s*(?!17\d{3})\d{4,5}(?:[\-\.\w]*)?(?:[^\n]*)', re.IGNORECASE),

            re.compile(r'EN\s*(?:IEC\s*)?(?!17\d{3})\d{2,5}(?:[\-\.\w]*)?(?:[^\n]*)', re.IGNORECASE),

            re.compile(r'IRAM\s*(?:NM\s*)?\d{3,5}(?:[\-\.\w]*)?(?:[^\n]*)', re.IGNORECASE),

            re.compile(r'NM\s*\d{3}(?:[\-\.\w]*)?', re.IGNORECASE),

            re.compile(r'ASTM\s*[A-Z]\d+(?:[\-\.\w]*)?', re.IGNORECASE),

        ]

        found = []

        seen_spans: set = set()

        for pattern in patterns:

            for m in pattern.finditer(text):

                span = m.span()

                if any(span[0] >= s[0] and span[1] <= s[1] for s in seen_spans):

                    continue

                norm = m.group(0).strip().rstrip(".,;")

                norm = re.sub(r' {2,}', ' ', norm)

                if not norm:

                    continue

                norm_l = norm.lower()

                if any(norm_l in f.lower() for f in found):

                    continue

                found = [f for f in found if f.lower() not in norm_l]

                found.append(norm)

                seen_spans.add(span)

        return ", ".join(found) if found else ""



    def _extract_date(self, text: str, date_type: str) -> str:

        if date_type == "emision":

            context_patterns = [r"(?:Date\s*of\s*[Ii]ssu(?:e|ance)|Fecha\s*de\s*[Ee]misi[oó]n|[Ii]ssued?)\s*:?\s*"]

        else:

            context_patterns = [r"(?:Valid\s*(?:until|to|through)|Expir[ey]|Vencimiento|[Vv]igencia\s*hasta)\s*:?\s*"]

        date_pattern = r"(\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4})"

        for ctx in context_patterns:

            match = re.search(ctx + date_pattern, text, re.IGNORECASE)

            if match:

                return match.group(1)

        return ""



    def _generate_djc_link(self, bidcom_num: str) -> str:

        # URL con /certificacion- y el número Bidcom (no el número del certificado OEC)

        base = "https://qr.gadnic.com/certifications/certificacion-"

        num_clean = str(bidcom_num).replace("C", "").replace("c", "").strip()

        if num_clean:

            return f"{base}{num_clean}"

        return ""



    def _infer_product_description(self, json_data: dict) -> str:

        tipo = json_data.get("tipo_intervencion", "")

        if tipo:

            return tipo

        specs = json_data.get("specs_tecnicas", "")

        if isinstance(specs, list) and specs:

            return specs[0][:50] if len(specs[0]) > 50 else specs[0]

        elif isinstance(specs, str) and specs:

            return str(specs)[:50]

        return ""



    def calculate_dates(self, fecha_emision: str, fecha_vencimiento: str = "") -> dict:

        result = {

            "fecha_inicio": fecha_emision,

            "fecha_vencimiento": fecha_vencimiento,

            "fecha_inicio_tramite": "",

        }

        if not fecha_vencimiento and fecha_emision:

            try:

                fe = parse_date(fecha_emision)

                result["fecha_vencimiento"] = (fe + timedelta(days=730)).strftime("%d/%m/%Y")

            except ValueError:

                pass

        if result["fecha_vencimiento"]:

            try:

                fv = parse_date(result["fecha_vencimiento"])

                result["fecha_inicio_tramite"] = (fv - timedelta(days=90)).strftime("%d/%m/%Y")

            except ValueError:

                pass

        return result



    def get_reglamento_options(self) -> list:

        return self.config.get("reglamento_options", [])



    def get_esquema_options(self) -> list:

        return self.config.get("esquema_options", [])



    def get_oec_options(self) -> dict:

        return self.config.get("oec_options", {})





# ─────────────────────────────────────────────────────────────

#  CLI para testing rápido

# ─────────────────────────────────────────────────────────────



if __name__ == "__main__":

    import sys

    gen = DJCGenerator()

    if len(sys.argv) > 1:

        cert_path = sys.argv[1]

        print(f"Extrayendo datos de: {cert_path}")

        cert_data = gen.extract_cert_data(cert_path)

        print(f"\n  Cert Number: {cert_data['cert_number']}")

        print(f"  Normas:      {cert_data['normas']}")

        print(f"  Emisión:     {cert_data['fecha_emision']}")

        print(f"  Vencimiento: {cert_data['fecha_vencimiento']}")

        print(f"  OEC:         {cert_data['oec_key']}")

        reglamento = gen.detect_reglamento(cert_data["normas"])

        print(f"  Reglamento:  {reglamento}")

    else:

        print("Uso: python m3_djc_generator.py <cert.pdf>")

        print("\nReglamentos disponibles:")

        for r in gen.get_reglamento_options():

            print(f"  • {r}")

