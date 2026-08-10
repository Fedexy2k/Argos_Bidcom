"""
Módulo 4: Generador de DJC de Eficiencia Energética (EE)
=========================================================
Encargado de construir las especificaciones técnicas dinámicas,
completar la plantilla oficial DJ Conformidad Modelo EE.docx,
e insertar la etiqueta oficial autogenerada como imagen.
"""

from __future__ import annotations

import io
import os
import json
import logging
import zoneinfo
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Inches, Pt, Cm  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

logger = logging.getLogger(__name__)

from modules.m3_djc_generator import normalize_oec_key

class DJCEEGenerator:
    """Genera Declaraciones Juradas de Conformidad de Eficiencia Energética (DJC-EE)."""

    TEMPLATE_FILENAME = "DJ Conformidad Modelo EE.docx"
    FT_TEMPLATE_FILENAME = "Ficha Tecnica Modelo EE.docx"

    def __init__(self, config_path: Optional[str] = None, ee_config_path: Optional[str] = None):
        base_dir = Path(__file__).parent.parent
        
        if config_path is None:
            config_path = str(base_dir / "m3_config.json")
        if ee_config_path is None:
            ee_config_path = str(base_dir / "ee_families.json")

        with open(config_path, "r", encoding="utf-8") as f:
            self.config = json.load(f)

        with open(ee_config_path, "r", encoding="utf-8") as f:
            self.ee_families = json.load(f)["families"]

        self.template_path = base_dir / self.TEMPLATE_FILENAME
        if not self.template_path.exists():
            # Fallback a ejemplos si no está en la raíz
            self.template_path = base_dir / "ejemplos" / "DJ Conformidad Modelo EE (1).docx"
            if not self.template_path.exists():
                raise FileNotFoundError(f"Template DJC-EE no encontrado en {base_dir / self.TEMPLATE_FILENAME}")

        self.ft_template_path = base_dir / self.FT_TEMPLATE_FILENAME

        logger.info(f"DJCEEGenerator inicializado. Template DJC: {self.template_path.name}")

    def get_family_by_id(self, family_id: str) -> Optional[dict]:
        """Busca una familia por su ID en el listado de familias."""
        for fam in self.ee_families:
            if fam["id"] == family_id:
                return fam
        return None

    def auto_extract_ee_from_report(self, report_text: str) -> Optional[dict]:
        """
        Usa la IA para extraer automáticamente la familia y los campos de Eficiencia Energética
        desde el texto de un informe de ensayo o certificado EE.
        """
        try:
            from modules.ai_helper import extract_ee_specs_ai
            res = extract_ee_specs_ai(report_text, self.ee_families)
            if not res:
                return None

            # Cálculo automático de Fecha de Próxima Vigilancia (+4 años desde emisión)
            fem = res.get("fecha_emision")
            if fem and not res.get("fecha_proxima_vigilancia"):
                try:
                    parts = fem.split('/')
                    if len(parts) == 3:
                        d, m, y = parts
                        res["fecha_proxima_vigilancia"] = f"{d}/{m}/{int(y)+4}"
                except Exception:
                    pass

            # Cálculos automáticos de métricas EE si faltan
            ee_fields = res.get("ee_fields", {})
            if res.get("family_id") == "lavavajillas":
                # Consumo anual de agua: 280 x consumo por ciclo
                if ee_fields.get("agua_ciclo") and not ee_fields.get("agua_anual"):
                    try:
                        ee_fields["agua_anual"] = round(float(ee_fields["agua_ciclo"]) * 280, 1)
                    except Exception:
                        pass
                # Consumo anual de energía si falta: 280 x consumo por ciclo
                if ee_fields.get("consumo_ciclo") and not ee_fields.get("consumo_anual"):
                    try:
                        ee_fields["consumo_anual"] = round(float(ee_fields["consumo_ciclo"]) * 280, 2)
                    except Exception:
                        pass

            return res
        except Exception as e:
            logger.error(f"Error en extracción automática de EE por IA: {e}")
            return None



    def build_specs_text(self, family_id: str, base_specs: dict, ee_fields_data: dict) -> str:
        """
        Compone el texto multilínea de especificaciones técnicas.
        Combina las specs eléctricas base y los campos dinámicos específicos.
        """
        lines = []

        # 1. Specs eléctricas base
        base_parts = []
        if base_specs.get("tension"):
            base_parts.append(base_specs["tension"])
        if base_specs.get("frecuencia"):
            base_parts.append(base_specs["frecuencia"])
        if base_specs.get("potencia"):
            base_parts.append(base_specs["potencia"])
        if base_specs.get("clase"):
            base_parts.append(base_specs["clase"])
        if base_specs.get("ip"):
            base_parts.append(base_specs["ip"])
        if base_specs.get("adicionales"):
            base_parts.append(base_specs["adicionales"])

        if base_parts:
            lines.append("; ".join(base_parts))

        # 2. Campos dinámicos de la familia
        family = self.get_family_by_id(family_id)
        if family:
            for field in family.get("fields", []):
                key = field["key"]
                label = field["label"]
                unit = field.get("unit", "")
                
                # Ignorar si es opcional y no se proveyó valor
                val = ee_fields_data.get(key)
                if val is None or str(val).strip() == "":
                    continue

                val_str = str(val).strip()
                
                # Formatear línea
                # Para evitar repetir la palabra consumo o clase de forma redundante:
                if key == "clase_ee":
                    lines.append(f"Clase de eficiencia energética: {val_str}")
                elif key == "clase_secado":
                    lines.append(f"Eficacia de secado: {val_str}")
                elif key == "capacidad":
                    lines.append(f"Capacidad declarada {val_str} {unit}")
                elif key == "ruido":
                    lines.append(f"Nivel de ruido: {val_str} {unit}")
                else:
                    unit_str = f" {unit}" if unit else ""
                    # Asegurar mayúscula en la etiqueta para legibilidad
                    lines.append(f"{label}: {val_str}{unit_str}")

        return "\n".join(lines)

    def generate_djc_id(self, bidcom_num: str, emision_date_str: str) -> str:
        """
        Genera el código ID de la DJC: DJC-EE-{MMYY}-C{NUM}-V1.
        """
        try:
            # Parsear la fecha de emisión del informe para obtener MMYY
            dt = datetime.strptime(emision_date_str, "%d/%m/%Y")
            anio_mes = dt.strftime("%m%y")
        except Exception:
            anio_mes = datetime.now().strftime("%m%y")

        bidcom_clean = str(bidcom_num).replace("/", "-").replace("C", "").replace("c", "") if bidcom_num else "XXXX"
        return f"DJC-EE-{anio_mes}-C{bidcom_clean}-V1"

    def fill_template_ee(self, data: dict, images_bytes: Optional[list] = None) -> Document:
        """Llena la plantilla oficial DJC-EE con los datos correspondientes.
        
        Estructura real del template (7 tablas, índices 0-6):
          Tabla 0: ID DJC (2 filas x 1 col)
          Tabla 1: Empresa (7 filas x 2 cols)
          Tabla 2: Representante Autorizado (3 filas x 2 cols)
          Tabla 3: Producto (7 filas x 2 cols)
          Tabla 4: Ensayo y Laboratorio (10 filas x 3 cols)
          Tabla 5: Enlace + 3 filas de etiquetas (6 slots: filas 1-3 × cols 0-1) + Fecha + Lugar
          Tabla 6: Firma (2 filas x 2 cols)
        
        images_bytes: lista de bytes PNG, uno por modelo. Se insertan en orden:
          slot 0 → fila 1, col 0
          slot 1 → fila 1, col 1
          slot 2 → fila 2, col 0
          slot 3 → fila 2, col 1
          slot 4 → fila 3, col 0
          slot 5 → fila 3, col 1
        """
        doc = Document(str(self.template_path))
        tables = doc.tables

        if len(tables) < 7:
            raise ValueError(f"Template de Eficiencia Energética inválido: esperaba 7 tablas, encontró {len(tables)}. "
                             f"Revisá que el archivo sea 'DJ Conformidad Modelo EE.docx' correcto.")

        # --- Tabla 0: ID DJC ---
        self._set_cell_id(tables[0], 1, 0, data.get("djc_id", ""))

        # --- Tabla 1: Empresa ---
        emp = data.get("empresa_override") or self.config["empresa"]
        self._set_cell(tables[1], 0, 1, emp.get("razon_social", ""))
        self._set_cell(tables[1], 1, 1, emp.get("cuit", ""))
        self._set_cell(tables[1], 2, 1, emp.get("marca_registrada", ""))
        self._set_cell(tables[1], 3, 1, emp.get("domicilio_legal", ""))
        self._set_cell(tables[1], 4, 1, emp.get("domicilio_deposito", ""))
        self._set_cell(tables[1], 5, 1, emp.get("telefono", ""))
        self._set_cell(tables[1], 6, 1, emp.get("email", ""))

        # --- Tabla 2: Representante Autorizado (siempre No Aplica en EE) ---
        self._set_cell(tables[2], 0, 1, "No Aplica")
        self._set_cell(tables[2], 1, 1, "No Aplica")
        self._set_cell(tables[2], 2, 1, "No Aplica")

        # --- Tabla 3: Producto ---
        self._set_cell(tables[3], 0, 1, "ver Modelo más abajo")
        self._set_cell(tables[3], 1, 1, "Información Restringida - Res. SIyC 237/2024 (China)")
        self._set_cell(tables[3], 2, 1, "Información Restringida - Res. SIyC 237/2024 (China)")
        self._set_cell(tables[3], 3, 1, data.get("producto_desc", ""))
        self._set_cell(tables[3], 4, 1, data.get("marca", ""))
        self._set_cell(tables[3], 5, 1, data.get("modelo", ""))
        self._set_cell(tables[3], 6, 1, data.get("specs", ""))

        # --- Tabla 4: Ensayo y Laboratorio (10 filas x 3 cols) ---
        # Resolución 438/2024 en ambas columnas de reglamento (fila 0)
        self._set_cell(tables[4], 0, 1, "Resolución 438/2024 (Eficiencia Energética)")
        self._set_cell(tables[4], 0, 2, "Resolución 438/2024 (Eficiencia Energética)")

        # Normas (fila 1)
        normas = data.get("normas", "")
        self._set_cell(tables[4], 1, 1, normas)
        self._set_cell(tables[4], 1, 2, normas)

        self._set_cell(tables[4], 2, 2, data.get("cert_number", ""))            # N° Ensayo
        self._set_cell(tables[4], 3, 2, "Eficiencia Energética")                 # Esquema fijo
        self._set_cell(tables[4], 4, 2, data.get("fecha_emision", ""))           # Fecha emisión informe
        self._set_cell(tables[4], 5, 2, "No aplicable")                          # Última vigilancia
        self._set_cell(tables[4], 6, 2, data.get("fecha_proxima_vigilancia", "")) # Próxima vigilancia

        # Laboratorio (fila 7) y Nombre OEC (fila 8) — mismo nombre en ambas
        lab_name = data.get("oec_nombre", "")
        self._set_cell(tables[4], 7, 2, lab_name)
        self._set_cell(tables[4], 8, 2, lab_name)

        # Datos de contacto del laboratorio (fila 9)
        contacto = data.get("oec_contacto", "").strip()
        if not contacto:
            # Fallback al config por compatibilidad si se usa oec_key
            oec_key = data.get("oec_key", "")
            normalized_oec_key = normalize_oec_key(oec_key)
            lab_info = self.config["oec_options"].get(normalized_oec_key, {})
            contacto = lab_info.get("contacto", "")
        self._set_cell(tables[4], 9, 2, contacto)

        # --- Tabla 5: Enlace + Grid de etiquetas (3 filas x 2 cols = 6 slots) + Fecha + Lugar ---
        # Fila 0: Enlace a la DJC en internet
        enlace = data.get("enlace_djc", "")
        if enlace:
            self._set_cell_hyperlink(tables[5], 0, 1, enlace, enlace)
        else:
            self._set_cell(tables[5], 0, 1, "")

        # Guardar referencias directas a Fecha y Lugar ANTES de eliminar filas
        # (los índices cambian si se eliminan filas; las referencias al XML sobreviven)
        fecha_cell_ref = tables[5].rows[4].cells[1]
        lugar_cell_ref = tables[5].rows[5].cells[1]

        # Parsear modelos del campo modelo (separados por coma)
        modelos = [m.strip() for m in data.get("modelo", "").split(",") if m.strip()]
        images_list = images_bytes or []
        n = len(images_list)

        # Tres filas de etiquetas: cada fila tiene dos slots (col 0 y col 1)
        # slot_a, slot_b = índices en images_list; t5_row = índice de fila en Tabla 5
        slot_rows = [(0, 1, 1), (2, 3, 2), (4, 5, 3)]

        # Procesar de abajo hacia arriba para que la eliminación no desplace índices
        for slot_a, slot_b, t5_row in reversed(slot_rows):
            has_a = slot_a < n
            has_b = slot_b < n

            if not has_a and not has_b:
                # Ninguna imagen en esta fila → eliminar la fila completa
                self._delete_table_row(tables[5], t5_row)
            else:
                # Slot A (col 0)
                if has_a:
                    model_a = modelos[slot_a] if slot_a < len(modelos) else ""
                    self._insert_label_in_cell(tables[5], t5_row, 0, images_list[slot_a], model_a)
                # Slot B (col 1)
                if has_b:
                    model_b = modelos[slot_b] if slot_b < len(modelos) else ""
                    self._insert_label_in_cell(tables[5], t5_row, 1, images_list[slot_b], model_b)
                elif has_a:
                    # Fila con una sola imagen → marcador en la celda vacía
                    self._set_cell(tables[5], t5_row, 1, "----")

        # Fijar Fecha y Lugar mediante referencias directas (índices ya no son confiables)
        tz = zoneinfo.ZoneInfo("America/Argentina/Buenos_Aires")
        fecha_hoy = datetime.now(tz).strftime("%d/%m/%Y")
        self._fill_cell_direct(fecha_cell_ref, data.get("fecha_emision_djc", fecha_hoy))
        self._fill_cell_direct(lugar_cell_ref, self.config["emision"]["lugar"])

        # --- Tabla 6: Firma ---
        self._set_cell(tables[6], 1, 1, self.config["firma"]["aclaracion"])

        logger.info(f"Plantilla DJC-EE completada para ID: {data.get('djc_id')}")
        return doc

    # --- Helpers de celdas Word ---

    def _set_cell(self, table, row: int, col: int, value: str):
        try:
            cell = table.rows[row].cells[col]
            if cell.paragraphs:
                para = cell.paragraphs[0]
                if para.runs:
                    run = para.runs[0]
                    run.text = str(value) if value else ""
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
                    for extra_run in para.runs[1:]:
                        extra_run.text = ""
                else:
                    para.clear()
                    run = para.add_run(str(value) if value else "")
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
            else:
                cell.text = str(value) if value else ""
        except (IndexError, AttributeError) as e:
            logger.warning(f"No se pudo escribir en celda [{row},{col}] de la tabla: {e}")

    def _set_cell_id(self, table, row: int, col: int, value: str):
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
        from docx.oxml.ns import qn as _qn  # type: ignore[import-untyped]
        from lxml import etree  # type: ignore[import-untyped]
        try:
            cell = table.rows[row].cells[col]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for shd in tcPr.findall(_qn('w:shd')):
                tcPr.remove(shd)
            shd = etree.SubElement(tcPr, _qn('w:shd'))
            shd.set(_qn('w:val'), 'clear')
            shd.set(_qn('w:color'), 'auto')
            shd.set(_qn('w:fill'), 'FFFFFF')

            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(value) if value else "")
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
        except Exception as e:
            logger.warning(f"No se pudo formatear celda ID [{row},{col}]: {e}")
            self._set_cell(table, row, col, value)

    def _set_cell_hyperlink(self, table, row: int, col: int, url: str, display_text: Optional[str] = None):
        try:
            from docx.oxml.shared import OxmlElement, qn  # type: ignore[import-untyped]
            cell = table.rows[row].cells[col]
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.clear()

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), para.part.relate_to(
                url,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            ))

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '16')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '16')
            rPr.append(szCs)

            new_run.append(rPr)
            new_run.text = display_text if display_text else url
            hyperlink.append(new_run)
            para._element.append(hyperlink)
        except Exception as e:
            logger.warning(f"No se pudo crear hipervínculo en celda [{row},{col}]: {e}")
            self._set_cell(table, row, col, url)

    def _insert_label_in_cell(self, table, row: int, col: int, image_bytes: bytes, model_name: str = ""):
        """Inserta una etiqueta EE en una celda: leyenda superior + imagen centrada."""
        try:
            cell = table.rows[row].cells[col]
            # Limpiar placeholder anterior
            for para in cell.paragraphs:
                para.clear()

            # Párrafo 1: leyenda "ETIQUETA (model_name)" en negrita pequeña
            legend_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            legend_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            legend_run = legend_para.add_run(
                f"ETIQUETA ({model_name})" if model_name else "ETIQUETA"
            )
            legend_run.bold = True
            legend_run.font.size = Pt(8)

            # Párrafo 2: imagen
            img_para = cell.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(io.BytesIO(image_bytes), width=Cm(6.0))

            logger.info(f"Etiqueta '{model_name}' insertada en slot [{row},{col}].")
        except Exception as e:
            logger.error(f"Error al insertar etiqueta en slot [{row},{col}]: {e}")

    def _delete_table_row(self, table, row_idx: int):
        """Elimina una fila de tabla mediante manipulación directa del XML."""
        tbl = table._tbl
        tr = table.rows[row_idx]._tr
        tbl.remove(tr)
        logger.debug(f"Fila {row_idx} de tabla eliminada.")

    def _fill_cell_direct(self, cell, value: str):
        """Llena una celda usando su referencia directa (seguro después de eliminar filas)."""
        for para in cell.paragraphs:
            para.clear()
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.add_run(value)

    # --- Export a PDF (Reutilizando la lógica robusta del M3) ---

    def export_to_pdf(self, docx_path: str, pdf_path: Optional[str] = None) -> str:
        """Convierte la DJC de Word a PDF de forma robusta."""
        if pdf_path is None:
            pdf_path = docx_path.replace(".docx", ".pdf")

        docx_abs = os.path.abspath(docx_path)
        pdf_abs  = os.path.abspath(pdf_path)

        # Intento 1: MS Word via COM
        try:
            import comtypes.client  # type: ignore[import-untyped]
            import comtypes
            wdFormatPDF = 17
            word = None
            doc  = None
            try:
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(docx_abs)
                doc.SaveAs(pdf_abs, FileFormat=wdFormatPDF)
                logger.info(f"DJC-EE PDF generado (via MS Word): {pdf_abs}")
                return pdf_abs
            finally:
                if doc is not None:
                    try:
                        doc.Close(False)
                    except Exception:
                        pass
                if word is not None:
                    try:
                        word.Quit()
                    except Exception:
                        pass
                    try:
                        del word
                    except Exception:
                        pass
        except Exception as e:
            logger.warning(f"docx2pdf (MS Word) falló o no está disponible: {e}. Intentando LibreOffice...")

        # Intento 2: LibreOffice
        import subprocess
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",
        ]

        soffice_exe = None
        for path in libreoffice_paths:
            if path == "soffice":
                try:
                    subprocess.run([path, "--version"], capture_output=True, check=True)
                    soffice_exe = path
                    break
                except Exception:
                    continue
            elif os.path.exists(path):
                soffice_exe = path
                break

        if soffice_exe:
            try:
                outdir = os.path.dirname(pdf_abs)
                subprocess.run(
                    [soffice_exe, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_abs],
                    check=True,
                    capture_output=True,
                )
                lo_output_path = os.path.join(outdir, os.path.splitext(os.path.basename(docx_abs))[0] + ".pdf")
                if lo_output_path != pdf_abs and os.path.exists(lo_output_path):
                    if os.path.exists(pdf_abs):
                        os.remove(pdf_abs)
                    os.rename(lo_output_path, pdf_abs)

                logger.info(f"DJC-EE PDF generado (via LibreOffice): {pdf_abs}")
                return pdf_abs
            except Exception as lo_err:
                logger.error(f"Error convirtiendo con LibreOffice: {lo_err}")
                raise RuntimeError(f"No se pudo generar PDF. Error LibreOffice: {lo_err}")
        else:
            logger.error("No se encontró Microsoft Word ni LibreOffice instalados.")
            raise RuntimeError(f"No se pudo generar PDF. LibreOffice no instalado.")

    def generate_ft_id(self, bidcom_num: str, emision_date_str: str) -> str:
        """Genera el código ID de la Ficha Técnica: FT-EE-{MMYY}-C{NUM}-V1."""
        try:
            dt = datetime.strptime(emision_date_str, "%d/%m/%Y")
            anio_mes = dt.strftime("%m%y")
        except Exception:
            anio_mes = datetime.now().strftime("%m%y")

        bidcom_clean = str(bidcom_num).replace("/", "-").replace("C", "").replace("c", "") if bidcom_num else "XXXX"
        return f"FT-EE-{anio_mes}-C{bidcom_clean}-V1"

    def fill_template_ft(self, data: dict, family_id: str, ee_fields_data: dict, images_bytes: Optional[list] = None) -> Document:
        """
        Llena la plantilla de Ficha Técnica de Información de Producto con los datos dinámicos.
        Filtra la Tabla 2 eliminando las filas que no pertenecen a la familia seleccionada.
        """
        import unicodedata

        def normalize_text(text: str) -> str:
            if not text:
                return ""
            text = unicodedata.normalize('NFD', text)
            text = "".join(c for c in text if unicodedata.category(c) != 'Mn')
            return text.lower().strip()

        if not self.ft_template_path.exists():
            raise FileNotFoundError(f"Template Ficha Técnica no encontrado en {self.ft_template_path}")

        doc = Document(str(self.ft_template_path))
        tables = doc.tables

        if len(tables) < 5:
            raise ValueError(f"Template Ficha Técnica inválido: esperaba 5 tablas, encontró {len(tables)}.")

        # --- Tabla 0: ID FT ---
        ft_id = data.get("ft_id") or data.get("djc_id", "").replace("DJC-EE-", "FT-EE-")
        if not ft_id:
            ft_id = self.generate_ft_id(data.get("bidcom_num", ""), data.get("fecha_emision", ""))
        self._set_cell_id(tables[0], 1, 0, ft_id)

        # --- Tabla 1: Identificación Comercial ---
        emp = data.get("empresa_override") or self.config["empresa"]
        marca = data.get("marca") or emp.get("marca_registrada", "BIDCOM")
        self._set_cell(tables[1], 0, 1, marca)
        self._set_cell(tables[1], 1, 1, data.get("modelo", ""))
        self._set_cell(tables[1], 2, 1, data.get("origen", "China"))

        # --- Tabla 2: Especificaciones Técnicas (Filtrar y Completar) ---
        family = self.get_family_by_id(family_id)
        ficha_fields = family.get("ficha_fields", []) if family else []
        norma_base = family.get("norma_base", "") if family else data.get("normas", "")

        label_map = {}
        static_map = {}
        for ff in ficha_fields:
            lbl_norm = normalize_text(ff["row_label"])
            if "data_key" in ff:
                label_map[lbl_norm] = ff["data_key"]
            elif "static_value" in ff:
                static_map[lbl_norm] = ff["static_value"]

        static_always_keep = {
            normalize_text("Norma Tecnica de Referencia"): norma_base,
            normalize_text("Resolucion Aplicable"): "Resolución SIyC N° 438/2024",
            normalize_text("Tipo / Categoria de Producto"): ee_fields_data.get("categoria", "") or (family.get("label", "") if family else ""),
        }

        table_specs = tables[2]
        rows_indices = list(range(len(table_specs.rows)))
        rows_indices.reverse()

        for idx in rows_indices:
            row = table_specs.rows[idx]
            cell_label = row.cells[0].text.strip()
            cell_norm = normalize_text(cell_label)

            is_static = False
            for stat_norm, stat_val in static_always_keep.items():
                if stat_norm in cell_norm or cell_norm in stat_norm:
                    is_static = True
                    if stat_val:
                        self._fill_cell_direct(row.cells[1], str(stat_val))
                    break

            if not is_static:
                for stat_norm, stat_val in static_map.items():
                    if stat_norm in cell_norm or cell_norm in stat_norm:
                        is_static = True
                        if stat_val:
                            self._fill_cell_direct(row.cells[1], str(stat_val))
                        break

            if is_static:
                continue

            matched_key = None
            for lbl_norm, key in label_map.items():
                if lbl_norm in cell_norm or cell_norm in lbl_norm:
                    matched_key = key
                    break

            if matched_key:
                val = ee_fields_data.get(matched_key, "")
                if val is not None and str(val).strip() != "":
                    unit = ""
                    if family:
                        for f_info in family.get("fields", []):
                            if f_info["key"] == matched_key and "unit" in f_info:
                                unit = f" {f_info['unit']}"
                                break
                    self._fill_cell_direct(row.cells[1], f"{val}{unit}")
                else:
                    self._fill_cell_direct(row.cells[1], "N/A")
            else:
                self._delete_table_row(table_specs, idx)

        # --- Tabla 3: Etiquetas y QR ---
        t3 = tables[3]
        qr_url = data.get("qr_url", f"https://qr.gadnic.com/certifications/certificado-{data.get('cert_number', 'ZZZ')}-ee")
        self._set_cell_hyperlink(t3, 0, 1, qr_url, qr_url)

        if images_bytes:
            model_names = [m.strip() for m in data.get("modelo", "").split("/") if m.strip()]
            slots = [(1, 0), (1, 1), (2, 0), (2, 1), (3, 0), (3, 1)]
            for idx_img, img_b in enumerate(images_bytes[:len(slots)]):
                r, c = slots[idx_img]
                m_name = model_names[idx_img] if idx_img < len(model_names) else f"Modelo {idx_img+1}"
                self._insert_label_in_cell(t3, r, c, img_b, model_name=m_name)

        fecha_emision = data.get("fecha_emision_djc") or datetime.now().strftime("%d/%m/%Y")
        self._set_cell(t3, 4, 1, fecha_emision)
        self._set_cell(t3, 5, 1, data.get("lugar", "Ciudad de Buenos Aires"))

        # --- Tabla 4: Firma ---
        aclaracion = data.get("firmante_override") or data.get("aclaracion", "BARNA, Emanuel Lucas")
        self._set_cell(tables[4], 1, 1, aclaracion)

        return doc
