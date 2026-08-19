"""
m5_solicitud_generator.py — Generador Automático de Solicitudes de Certificación
==================================================================================

Módulo principal para generar solicitudes de certificación (Lenor y qetkra)
a partir de planillas de ingeniería (datasheets).

Responsabilidades:
  1. parse_datasheet(filepath)            → extrae metadatos y modelos del Excel de ingeniería
  2. generate_lenor(data, svg_bytes, nro) → llena solicitud Lenor Excel + Nota Word + PDF QR
  3. generate_qetkra(data, nro)           → llena solicitud qetkra Excel + Nota Word
  4. save_and_zip(nro, files)             → guarda en Solicitudes/[Nro]/ y devuelve ZIP bytes
"""
from __future__ import annotations

import io
import os
import re
import unicodedata
import shutil
import zipfile
from datetime import date
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document
from copy import deepcopy
from modules.regulations import suggest_reg_and_norm

# Carga segura de win32com
try:
    import win32com.client
    import pythoncom
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

# ── Rutas base ────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = ROOT / "assets" / "solicitud_templates"
OUTPUT_BASE = ROOT / "Solicitudes"

TEMPLATE_LENOR_XLSX = TEMPLATES_DIR / "Solicitud_Modelo_Lenor.xlsm"
TEMPLATE_LENOR_DOCX = TEMPLATES_DIR / "Nota_Modelo_Lenor.docx"
TEMPLATE_qetkra_XLSX = TEMPLATES_DIR / "Solicitud_Modelo_qetkra.xlsx"
TEMPLATE_qetkra_DOCX = TEMPLATES_DIR / "Nota_Modelo_qetkra.docx"
TEMPLATE_TUV_DOCX = TEMPLATES_DIR / "Solicitud_Modelo_tuv.docx"


# ═════════════════════════════════════════════════════════════════════════════
# 1. LOGGING HELPER & PARSER DE DATASHEET
# ═════════════════════════════════════════════════════════════════════════════

def _log(msg: str, level: str = "info", logger: Any = None) -> None:
    """Helper de logging que envía logs a un logger provisto o a stdout."""
    if logger is not None:
        try:
            if hasattr(logger, "log"):
                logger.log(msg, level.upper())
            elif level.lower() == "error" and hasattr(logger, "error"):
                logger.error(msg)
            elif level.lower() == "warning" and hasattr(logger, "warning"):
                logger.warning(msg)
            elif level.lower() == "debug" and hasattr(logger, "debug"):
                logger.debug(msg)
            elif hasattr(logger, "info"):
                logger.info(msg)
        except Exception:
            print(f"[{level.upper()}] {msg}")
    else:
        print(f"[{level.upper()}] {msg}")


def _clean_key(val: Any) -> str:
    """Normaliza un valor de celda a clave de búsqueda (minúsculas, sin acentos ni puntuación extra)."""
    if val is None:
        return ""
    s = str(val).lower().strip()
    s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
    s = s.replace(":", "").replace(".", "")
    s = re.sub(r"\s+", " ", s)
    return s


def _detect_oec(metadata: dict) -> str:
    """
    Detecta la certificadora a partir del motivo y el campo OEC del datasheet.
    Retorna 'tuv', 'qetkra' o 'lenor'.
    """
    motivo = metadata.get("motivo", "").lower()
    oec = metadata.get("oec", "").lower()
    if "tüv" in oec or "tuv" in oec or "tüv" in motivo or "tuv" in motivo:
        return "tuv"
    if "convenio" in motivo or "quektra" in oec or "qetkra" in oec:
        return "qetkra"
    return "lenor"


def split_marcas(marca_str: str) -> list[str]:
    """
    Separa las marcas por barra inclinada (/), punto y coma (;), coma (,) o la conjunción "o".
    Ejemplo: 'GADNIC; CARE BY GADNIC' -> ['GADNIC', 'CARE BY GADNIC']
    """
    if not marca_str:
        return []
    parts = re.split(r'\s*[/;,]\s*|\s+\bo\b\s+', marca_str, flags=re.IGNORECASE)
    return [p.strip() for p in parts if p.strip()]




def parse_specs_string(specs_str: str) -> dict:
    """
    Parsea una cadena de especificaciones técnicas y la divide en un diccionario
    con claves estructuradas para el anexo de modelos de Qetkra.
    """
    res = {
        "tension": "",
        "frecuencia": "",
        "corriente": "",
        "potencia": "",
        "aislacion": "",
        "tension_salida": "",
        "corriente_salida": "",
        "grado_ip": "",
        "casquillo": "",
        "adicional": ""
    }
    if not specs_str:
        return res

    # Dividir primero por ';' o saltos de línea
    parts = []
    for part in re.split(r'[;\n]', specs_str):
        part = part.strip()
        if part:
            # Separar por comas si no contiene caracteres descriptivos complejos
            subparts = re.split(r',', part)
            for sp in subparts:
                sp = sp.strip()
                if sp:
                    parts.append(sp)

    leftovers = []
    
    aislacion_pattern = r'\b(clase\s+[i|v|x]+|class\s+[i|v|x]+|\bclase\s+[1-3]\b|\bclass\s+[1-3]\b)\b'
    ip_pattern = r'\bip\s*[x0-9]{2,4}\b'
    freq_pattern = r'\b\d{2,3}(?:\s*[-/]\s*\d{2,3})?\s*hz\b'
    casquillo_pattern = r'\b(e14|e27|gu10|mr16|b22|e40)\b'

    for part in parts:
        part_lower = part.lower()
        matched_any = False
        
        # Clase de aislación
        m_ais = re.search(aislacion_pattern, part_lower)
        if m_ais:
            res["aislacion"] = m_ais.group(1).replace("class", "clase").title()
            # Normalizar números romanos
            res["aislacion"] = re.sub(r'\bclase\s+ii\b|\bclase\s+2\b', 'Clase II', res["aislacion"], flags=re.IGNORECASE)
            res["aislacion"] = re.sub(r'\bclase\s+iii\b|\bclase\s+3\b', 'Clase III', res["aislacion"], flags=re.IGNORECASE)
            res["aislacion"] = re.sub(r'\bclase\s+i\b|\bclase\s+1\b', 'Clase I', res["aislacion"], flags=re.IGNORECASE)
            part = re.sub(aislacion_pattern, "", part, flags=re.IGNORECASE).strip()
            part = re.sub(r'^[\s,;-]+|[\s,;-]+$', '', part)
            matched_any = True
            if not part:
                continue
            part_lower = part.lower()

        # Grado IP
        m_ip = re.search(ip_pattern, part_lower)
        if m_ip:
            res["grado_ip"] = m_ip.group(0).upper().replace(" ", "")
            part = re.sub(ip_pattern, "", part, flags=re.IGNORECASE).strip()
            part = re.sub(r'^[\s,;-]+|[\s,;-]+$', '', part)
            matched_any = True
            if not part:
                continue
            part_lower = part.lower()

        # Frecuencia
        m_freq = re.search(freq_pattern, part_lower)
        if m_freq:
            res["frecuencia"] = m_freq.group(0).strip()
            res["frecuencia"] = re.sub(r'\s*hz', 'Hz', res["frecuencia"], flags=re.IGNORECASE)
            part = re.sub(freq_pattern, "", part, flags=re.IGNORECASE).strip()
            part = re.sub(r'^[\s,;-]+|[\s,;-]+$', '', part)
            matched_any = True
            if not part:
                continue
            part_lower = part.lower()

        # Casquillo
        m_casq = re.search(casquillo_pattern, part_lower)
        if m_casq:
            res["casquillo"] = m_casq.group(1).upper()
            part = re.sub(casquillo_pattern, "", part, flags=re.IGNORECASE).strip()
            part = re.sub(r'^[\s,;-]+|[\s,;-]+$', '', part)
            matched_any = True
            if not part:
                continue
            part_lower = part.lower()

        # Voltajes, corrientes y potencias
        voltages = re.findall(r'\b\d+(?:\s*[-/]\s*\d+)?\s*v[~a-zA-Z]*', part, re.IGNORECASE)
        currents = re.findall(r'\b\d+(?:[.,]\d+)?\s*a\b', part, re.IGNORECASE)
        powers = re.findall(r'\b\d+(?:[.,]\d+)?\s*w\b', part, re.IGNORECASE)

        is_output = "salida" in part_lower or "output" in part_lower or "out" in part_lower or "vcc" in part_lower or "vd.c." in part_lower or "vdc" in part_lower
        is_input = "entrada" in part_lower or "input" in part_lower or "in" in part_lower or "v~" in part_lower or "vca" in part_lower or "va.c." in part_lower

        if voltages:
            matched_any = True
            for v in voltages:
                v_clean = v.strip()
                v_lower = v_clean.lower()
                if "vcc" in v_lower or "vdc" in v_lower or "vd.c." in v_lower or is_output:
                    if not res["tension_salida"]:
                        res["tension_salida"] = v_clean
                else:
                    if not res["tension"]:
                        res["tension"] = v_clean

        if currents:
            matched_any = True
            for c in currents:
                c_clean = c.strip()
                if is_output or (res["tension_salida"] and not is_input):
                    if not res["corriente_salida"]:
                        res["corriente_salida"] = c_clean
                else:
                    if not res["corriente"]:
                        res["corriente"] = c_clean

        if powers:
            matched_any = True
            p_val = part.strip()
            if res["potencia"]:
                if p_val not in res["potencia"]:
                    res["potencia"] += " / " + p_val
            else:
                res["potencia"] = p_val
            continue

        if not matched_any:
            clean_part = part.strip()
            clean_part = re.sub(r'^[\s,;-]+|[\s,;-]+$', '', clean_part)
            if len(clean_part) > 2 and clean_part.lower() not in ("especificaciones", "especificaciones:", "especificaciones técnicas"):
                leftovers.append(clean_part)

    if leftovers:
        res["adicional"] = "; ".join(leftovers)
        
    return res


def parse_datasheet(filepath: str | Path, logger: Any = None) -> dict:
    """
    Parsea una planilla de ingeniería Excel y retorna un diccionario estructurado.
    Soporta dos formatos:
      - Vertical (bloques SKU): ESTUPS, etc.  → lista de bloques con modelos por SKU
      - Tabular (tabla al final):  AIRCON3X, etc. → tabla con columnas SKU/Marca/Modelo/CB/Specs
    """
    _log(f"Iniciando parseo de datasheet: {filepath}", "info", logger)
    try:
        wb = openpyxl.load_workbook(str(filepath), data_only=True)
        _log("Archivo Excel cargado correctamente en memoria con openpyxl.", "info", logger)
    except Exception as e:
        _log(f"Error al abrir el Excel {filepath}: {e}", "error", logger)
        raise

    ws = wb.active
    _log(f"Hoja activa detectada: '{ws.title}' con dimensiones {ws.max_row}x{ws.max_column}", "info", logger)

    metadata: dict[str, str] = {}
    skus: list[dict] = []
    has_local_cert = False

    # ── Paso 1: buscar tabla con encabezado SKU/Modelo ────────────────────────
    table_header_row: int | None = None
    col_map: dict[str, int] = {}

    _log("Buscando encabezados de tabla (SKU, Modelo, etc.)...", "info", logger)
    for row in range(1, ws.max_row + 1):
        row_vals = [ws.cell(row, c).value for c in range(1, ws.max_column + 1)]
        row_clean = [_clean_key(v) for v in row_vals]

        if "sku" in row_clean and ("modelo" in row_clean or "modelos" in row_clean):
            table_header_row = row
            # Mapear columnas por nombre
            for ci, h in enumerate(row_clean, 1):
                if h == "sku":
                    col_map["sku"] = ci
                elif h == "marca":
                    col_map["marca"] = ci
                elif h in ("modelo", "modelos") and "alternativo" not in h and "cb" not in h and "fabrica" not in h and "certificado" not in h:
                    col_map["modelo"] = ci
                elif any(k in h for k in ("cb", "alternativo", "certificado", "fabrica")) and h != "fabrica":
                    col_map["modelo_fabrica"] = ci
                elif any(k in h for k in ("especificacion", "tecnica")):
                    col_map["specs"] = ci
                elif "tension" in h:
                    col_map["tension"] = ci
                elif "frecuencia" in h:
                    col_map["frecuencia"] = ci
                elif "corriente" in h:
                    col_map["corriente"] = ci
                elif "potencia" in h:
                    col_map["potencia"] = ci
                elif "aislacion" in h or "aislación" in h:
                    col_map["aislacion"] = ci
            _log(f"Tabla de modelos encontrada en fila {row}. Mapeo de columnas: {col_map}", "info", logger)
            break

    # Leer filas de tabla si existe
    if table_header_row and col_map.get("sku"):
        _log(f"Leyendo filas de la tabla de modelos desde fila {table_header_row + 1}...", "info", logger)
        current_sku: dict | None = None
        for r in range(table_header_row + 1, ws.max_row + 1):
            def _cv(col_key: str) -> str:
                ci = col_map.get(col_key)
                if not ci:
                    return ""
                v = ws.cell(r, ci).value
                return str(v).strip() if v is not None else ""

            sku_val = _cv("sku")
            marca_val = _cv("marca")
            modelo_val = _cv("modelo")
            modelo_fab = _cv("modelo_fabrica")

            if not sku_val and not modelo_val:
                continue

            # Agrupar modelos por SKU
            if current_sku is None or current_sku["sku"] != sku_val:
                if current_sku is not None:
                    _log(f"Bloque SKU finalizado: {current_sku['sku']} con marcas {split_marcas(current_sku['marca'])} y modelos {current_sku['modelos']}", "info", logger)
                    skus.append(current_sku)
                current_sku = {
                    "sku": sku_val,
                    "marca": marca_val,
                    "modelos": [],
                    "modelo_fabrica": modelo_fab,
                    "tension": _cv("tension"),
                    "frecuencia": _cv("frecuencia"),
                    "corriente": _cv("corriente"),
                    "potencia": _cv("potencia"),
                    "aislacion": _cv("aislacion"),
                    "specs": _cv("specs"),
                }
                _log(f"Nuevo bloque SKU detectado en fila {r}: SKU={sku_val}, Marca={marca_val}", "info", logger)
            if modelo_val:
                current_sku["modelos"].append(modelo_val)
            if modelo_fab and not current_sku["modelo_fabrica"]:
                current_sku["modelo_fabrica"] = modelo_fab

        if current_sku:
            _log(f"Último bloque SKU finalizado: {current_sku['sku']} con marcas {split_marcas(current_sku['marca'])} y modelos {current_sku['modelos']}", "info", logger)
            skus.append(current_sku)
    else:
        _log("No se detectó tabla tabular con encabezado SKU/Modelo. Se buscarán bloques verticales.", "info", logger)

    # ── Paso 2: extraer metadatos clave-valor ──
    max_meta_row = (table_header_row - 1) if table_header_row else ws.max_row
    _log(f"Buscando metadatos clave-valor hasta la fila {max_meta_row}...", "info", logger)

    vertical_blocks: list[dict] = []
    current_block: dict = {}

    for row in range(1, max_meta_row + 1):
        for col in range(1, ws.max_column + 1):
            cell_val = ws.cell(row, col).value
            if cell_val is None:
                continue
            key = _clean_key(cell_val)
            next_val_raw = ws.cell(row, col + 1).value
            next_val = str(next_val_raw).strip() if next_val_raw is not None else ""

            if key in (
                "nº de certificado local", "nº de certificado local:",
                "n° de certificado local", "n° de certificado local:",
                "certificado local", "certificado local:"
            ):
                val = next_val
                m = re.match(r"certificado\s+(\d+)", val, re.IGNORECASE)
                if m:
                    val = "C" + m.group(1)
                metadata["certificado"] = val
                has_local_cert = True
                _log(f"Metadato extraído: Certificado local='{val}'", "info", logger)

            elif key in (
                "nº de certificado", "nº de certificado:",
                "certificado", "certificado:",
                "n° de certificado", "n° de certificado:"
            ):
                val = next_val
                m = re.match(r"certificado\s+(\d+)", val, re.IGNORECASE)
                if m:
                    val = "C" + m.group(1)
                    metadata["certificado"] = val
                    has_local_cert = True
                    _log(f"Metadato extraído: Certificado local (desde certificado)='{val}'", "info", logger)
                else:
                    if has_local_cert:
                        metadata["certificado_origen"] = val
                        _log(f"Metadato extraído: Certificado de origen='{val}'", "info", logger)
                    else:
                        metadata["certificado"] = val
                        _log(f"Metadato extraído: Certificado='{val}'", "info", logger)

            elif key in ("motivo", "tipo de intervención", "tipo de intervencion"):
                metadata["motivo"] = next_val
                _log(f"Metadato extraído: Motivo='{next_val}'", "info", logger)

            elif key in (
                "organismo certificador", "organismo", "organismo certificador:"
            ):
                metadata["oec"] = next_val
                _log(f"Metadato extraído: OEC='{next_val}'", "info", logger)

            elif key == "normas":
                metadata["normas"] = next_val
                _log(f"Metadato extraído: Normas='{next_val}'", "info", logger)

            elif key in ("marca", "marca:"):
                metadata["marca"] = next_val
                _log(f"Metadato extraído: Marca='{next_val}'", "info", logger)

            elif key in ("especificaciones", "specs", "especificaciones:", "especificaciones técnicas"):
                metadata["specs"] = next_val
                _log(f"Metadato extraído: Especificaciones (global)='{next_val}'", "info", logger)

            elif key in ("producto", "descripción", "descripcion"):
                metadata["producto"] = next_val
                _log(f"Metadato extraído: Producto/Denominación='{next_val}'", "info", logger)

            elif key == "laboratorio":
                metadata["laboratorio"] = next_val
                _log(f"Metadato extraído: Laboratorio='{next_val}'", "info", logger)

            elif key == "reglamento":
                metadata["reglamento"] = next_val
                _log(f"Metadato extraído: Reglamento='{next_val}'", "info", logger)

            elif key in (
                "fábrica", "fabrica", "fabricante", "titular del certificado",
                "fábrica:", "fabricante:",
            ):
                metadata["fabrica"] = next_val
                _log(f"Metadato extraído: Fábrica/Fabricante='{next_val}'", "info", logger)

            elif key in (
                "dirección", "direccion", "dirección fábrica", "direccion fabrica",
                "dirección:", "dirección fábrica:",
            ):
                metadata["direccion"] = next_val
                _log(f"Metadato extraído: Dirección='{next_val}'", "info", logger)

            elif key in ("contacto", "contacto:"):
                metadata["contacto"] = next_val
                _log(f"Metadato extraído: Contacto='{next_val}'", "info", logger)

            elif key in (
                "email", "email proveedor", "e-mail", "email proveedor:",
                "email:", "e-mail:",
            ):
                metadata["email"] = next_val
                _log(f"Metadato extraído: Email='{next_val}'", "info", logger)

            elif (
                key.startswith("teléfono") or key.startswith("telefono")
                or key.startswith("teléfono /") or key.startswith("telefono /")
            ):
                tel = next_val.rstrip(".0") if "." in next_val and next_val.replace(".", "").isdigit() else next_val
                metadata["telefono"] = tel
                _log(f"Metadato extraído: Teléfono='{tel}'", "info", logger)

            # Bloques verticales (ESTUPS)
            if not table_header_row:
                if key == "sku":
                    if current_block:
                        _log(f"Bloque vertical finalizado: SKU={current_block['sku']}, Marca={current_block.get('marca')}, Modelos={current_block.get('modelos')}", "info", logger)
                        vertical_blocks.append(current_block)
                    current_block = {
                        "sku": next_val, "modelos": [],
                        "marca": "", "modelo_fabrica": "",
                        "tension": "", "frecuencia": "", "corriente": "",
                        "potencia": "", "aislacion": "", "specs": "",
                    }
                    _log(f"Nuevo bloque vertical detectado: SKU={next_val}", "info", logger)
                elif key in ("marca", "marca:"):
                    current_block["marca"] = next_val
                elif key in (
                    "especificaciones", "specs", "especificaciones:", "especificaciones técnicas"
                ):
                    current_block["specs"] = next_val
                elif key in ("tension", "tensión", "tension:", "tensión:"):
                    current_block["tension"] = next_val
                elif key in ("frecuencia", "frecuencia nominal", "frecuencia:"):
                    current_block["frecuencia"] = next_val
                elif key in ("corriente", "intensidad", "corriente:"):
                    current_block["corriente"] = next_val
                elif key in ("potencia", "consumo", "potencia:"):
                    current_block["potencia"] = next_val
                elif key in ("aislacion", "aislación", "clase de aislacion", "aislacion:"):
                    current_block["aislacion"] = next_val
                elif key in ("fábrica", "fabrica") and not metadata.get("fabrica"):
                    current_block["fabrica"] = next_val
                elif key in ("dirección", "direccion") and not metadata.get("direccion"):
                    current_block["direccion"] = next_val
                elif key in ("contacto",) and not metadata.get("contacto"):
                    current_block["contacto"] = next_val
                elif (
                    key in ("modelo", "modelo bidcom", "modelo bidcom principal")
                    or re.match(r"modelo\s+\d+", key)
                    or key.startswith("modelo ")
                ):
                    if next_val:
                        current_block.setdefault("modelos", []).append(next_val)

    if not table_header_row and current_block:
        _log(f"Último bloque vertical finalizado: SKU={current_block['sku']}, Marca={current_block.get('marca')}, Modelos={current_block.get('modelos')}", "info", logger)
        vertical_blocks.append(current_block)

    if vertical_blocks:
        _log(f"Consolidando metadatos de {len(vertical_blocks)} bloques verticales...", "info", logger)
        for b in vertical_blocks:
            for k in ("fabrica", "direccion", "contacto", "email", "telefono"):
                if b.get(k) and not metadata.get(k):
                    metadata[k] = b[k]

    if not skus and vertical_blocks:
        skus = vertical_blocks

    wb.close()

    global_marca = metadata.get("marca", "")
    global_specs = metadata.get("specs", "")

    # Parse specs for SKU blocks if they are empty
    for sku_block in skus:
        # Fallback para marca si está vacía
        if not sku_block.get("marca") and global_marca:
            sku_block["marca"] = global_marca
            _log(f"Asignando marca global '{global_marca}' al SKU '{sku_block['sku']}'", "info", logger)

        specs_str = sku_block.get("specs", "")
        if not specs_str and global_specs:
            specs_str = global_specs
            sku_block["specs"] = global_specs
            _log(f"Asignando especificaciones globales '{global_specs}' al SKU '{sku_block['sku']}'", "info", logger)

        parsed = parse_specs_string(specs_str)
        
        # Only overwrite if they are empty in the sku_block
        for k in ("tension", "frecuencia", "corriente", "potencia", "aislacion"):
            if not sku_block.get(k):
                sku_block[k] = parsed[k]
                
        # Always set additional fields
        sku_block["tension_salida"] = parsed["tension_salida"]
        sku_block["corriente_salida"] = parsed["corriente_salida"]
        sku_block["grado_ip"] = parsed["grado_ip"]
        sku_block["casquillo"] = parsed["casquillo"]
        sku_block["adicional"] = parsed["adicional"]

    oec_det = _detect_oec(metadata)
    _log(f"Detección de OEC: '{oec_det}'", "info", logger)
    _log(f"Parseo de datasheet completo. Total SKUs: {len(skus)}", "info", logger)

    normas_val = metadata.get("normas", "")
    reglamento_val = metadata.get("reglamento", "")

    if (not normas_val or not reglamento_val) and oec_det in ("lenor", "qetkra"):
        sug = suggest_reg_and_norm(metadata.get("producto", ""))
        if not normas_val and sug["norma"]:
            normas_val = sug["norma"]
            _log(f"Campo 'normas' vacío en datasheet. Sugiriendo '{normas_val}' para producto '{metadata.get('producto', '')}' (OEC: {oec_det}).", "info", logger)
        if not reglamento_val and sug["reglamento"]:
            reglamento_val = sug["reglamento"]
            _log(f"Campo 'reglamento' vacío en datasheet. Sugiriendo '{reglamento_val}' para producto '{metadata.get('producto', '')}' (OEC: {oec_det}).", "info", logger)

    if not normas_val:
        _log("Campo 'normas' vacío en datasheet y sin sugerencia automática disponible.", "warning", logger)
    if not reglamento_val:
        _log("Campo 'reglamento' vacío en datasheet y sin sugerencia automática disponible.", "warning", logger)

    return {
        "oec_detected": oec_det,
        "certificado": metadata.get("certificado", ""),
        "certificado_origen": metadata.get("certificado_origen", ""),
        "producto": metadata.get("producto", ""),
        "motivo": metadata.get("motivo", ""),
        "oec": metadata.get("oec", ""),
        "normas": normas_val,
        "laboratorio": metadata.get("laboratorio", ""),
        "reglamento": reglamento_val,
        "fabrica": metadata.get("fabrica", ""),
        "direccion": metadata.get("direccion", ""),
        "contacto": metadata.get("contacto", ""),
        "email": metadata.get("email", ""),
        "telefono": metadata.get("telefono", ""),
        "skus": skus,
    }


# ═════════════════════════════════════════════════════════════════════════════
# 2. AUXILIARES DE FECHAS Y TEXTO
# ═════════════════════════════════════════════════════════════════════════════

def _today_ddmmyyyy() -> str:
    """Retorna la fecha de hoy en formato DD/MM/YYYY."""
    d = date.today()
    return d.strftime("%d/%m/%Y")


def _today_long_es() -> str:
    """Retorna la fecha de hoy en formato largo español (sin la ciudad)."""
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
    }
    d = date.today()
    return f"{d.day} de {meses[d.month]} de {d.year}"


def _get_base_standard(norm: str) -> str:
    n = norm.strip()
    m = re.match(r'^([a-zA-Z\s]+[\d]+(?:-[\d]+)*)', n)
    if m:
        return re.sub(r'\s+', ' ', m.group(1).strip()).upper()
    return n.upper()


def _split_normas(normas_str: str) -> list[str]:
    """Divide la cadena de normas en lista. La norma con '-1' va primero.
    Agrupa enmiendas (AMD, AMND, A1, etc.) con su norma base (ej: IEC 60335-2-65).
    """
    if not normas_str:
        return []
    
    # 1. Limpiar error común de coma entre la sigla y el número (ej: "IEC, 60335" -> "IEC 60335")
    normas_clean = re.sub(r'\b(IEC|IRAM|ISO|NM)\s*,\s*', r'\1 ', normas_str, flags=re.IGNORECASE)
    
    # 2. Dividir por coma, punto y coma, barra vertical, saltos de línea o guión rodeado de espacios
    raw_parts = [p.strip() for p in re.split(r'[,;|\n]|\s+-\s+', normas_clean) if p.strip()]
    
    # 3. Agrupar por norma base
    groups = {}
    order = []  # Para mantener el orden de aparición original
    
    for part in raw_parts:
        base_prefix = _get_base_standard(part)
        if base_prefix not in groups:
            groups[base_prefix] = []
            order.append(base_prefix)
        groups[base_prefix].append(part)
        
    consolidated = []
    for base_prefix in order:
        parts = groups[base_prefix]
        
        # Encontrar la parte base (la que no contiene enmiendas como AMD, AMND, A1, etc. o la más corta)
        base_candidates = [p for p in parts if not re.search(r'\b(AMD|AMND|A\d)\b', p, re.IGNORECASE)]
        
        if base_candidates:
            base_standard = min(base_candidates, key=len)
        else:
            base_standard = min(parts, key=len)
            
        # Extraer las enmiendas
        amendments = []
        for part in parts:
            if part == base_standard:
                continue
            # Quitar la parte base de la cadena para obtener solo la enmienda
            amend = part.replace(base_standard, "").strip("/ \t+,")
            if amend:
                amendments.append(amend)
                
        # Limpiar y ordenar enmiendas
        amendments = sorted(list(set(amendments)))
        
        # Combinar
        if amendments:
            combined = f"{base_standard} + " + " + ".join(amendments)
        else:
            combined = base_standard
            
        consolidated.append(combined)
        
    # Reordenar para que la norma "-1" vaya primero si existe
    main = [n for n in consolidated if "-1" in n]
    rest = [n for n in consolidated if "-1" not in n]
    return main + rest


# ═════════════════════════════════════════════════════════════════════════════
# 3. GENERADOR LENOR EXCEL (win32com + fallback openpyxl)
# ═════════════════════════════════════════════════════════════════════════════

def generate_lenor_excel_win32(data: dict, out_path: Path, logger: Any = None) -> bool:
    """Escribe la solicitud Lenor usando win32com para preservar macros y validaciones."""
    if not WIN32_AVAILABLE:
        _log("win32com no está instalado/disponible. Saltando a openpyxl.", "warning", logger)
        return False

    _log(f"Abriendo Excel (win32com) para Lenor. Destino: {out_path.name}", "info", logger)
    _log("Inicializando comunicación COM (pythoncom.CoInitialize)...", "debug", logger)
    pythoncom.CoInitialize()
    excel = None
    wb = None
    try:
        _log("Lanzando instancia de Excel en segundo plano (win32com.client.DispatchEx)...", "debug", logger)
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        _log("Instancia de Excel iniciada correctamente (Visible=False, DisplayAlerts=False).", "debug", logger)

        # Copiar plantilla al destino
        _log(f"Copiando plantilla {TEMPLATE_LENOR_XLSX.name} a {out_path.name}...", "info", logger)
        shutil.copy2(str(TEMPLATE_LENOR_XLSX), str(out_path))

        abs_out = str(out_path.resolve())
        _log(f"Abriendo archivo Excel destino: {abs_out}...", "debug", logger)
        wb = excel.Workbooks.Open(abs_out)

        # Hoja 1: Datos del solicitante y esquema
        _log("Buscando hoja de datos del solicitante (buscando 'datos' en el nombre de las hojas)...", "debug", logger)
        ws1 = None
        for s in wb.Worksheets:
            if "datos" in s.Name.lower():
                ws1 = s
                break
        if not ws1:
            _log("Hoja con 'datos' no encontrada, utilizando primera hoja por defecto.", "warning", logger)
            ws1 = wb.Worksheets(1)
        _log(f"Hoja de datos seleccionada: '{ws1.Name}'", "debug", logger)

        _log("Escribiendo metadatos principales en hoja de datos...", "info", logger)
        ws1.Range("G6").Value = "'" + _today_ddmmyyyy()
        _log(f"  - Fecha hoy en G6: '{_today_ddmmyyyy()}'", "debug", logger)



        fabrica = data.get("fabrica", "") or "----"
        direccion = data.get("direccion", "") or "----"
        contacto = data.get("contacto", "") or "----"
        email = data.get("email", "") or ""
        telefono = data.get("telefono", "") or ""

        ws1.Range("C40").Value = fabrica
        ws1.Range("C42").Value = direccion  # Domicilio legal
        ws1.Range("C44").Value = direccion  # Domicilio físico
        ws1.Range("C46").Value = contacto
        ws1.Range("E46").Value = email
        ws1.Range("G46").Value = telefono

        _log(f"  - Fábrica (C40): '{fabrica}'", "debug", logger)
        _log(f"  - Domicilio Legal (C42) y Físico (C44): '{direccion}'", "debug", logger)
        _log(f"  - Contacto (C46): '{contacto}', Email (E46): '{email}', Teléfono (G46): '{telefono}'", "debug", logger)

        normas_list = _split_normas(data.get("normas", ""))
        reglamento_val = data.get("reglamento", "") or ""
        norma_val = normas_list[0] if normas_list else ""
        ws1.Range("C57").Value = reglamento_val
        ws1.Range("C59").Value = norma_val
        _log(f"  - Reglamento (C57): '{reglamento_val}'", "debug", logger)
        _log(f"  - Norma principal (C59): '{norma_val}'", "debug", logger)

        # Hoja 2: Detalle de productos
        _log("Buscando hoja de detalle de productos (buscando 'detalle' en el nombre)...", "debug", logger)
        ws2 = None
        for s in wb.Worksheets:
            if "detalle" in s.Name.lower():
                ws2 = s
                break

        if ws2:
            _log(f"Hoja de detalle seleccionada: '{ws2.Name}'", "debug", logger)
            _log("Limpiando celdas del rango B11:F200 para evitar solapamientos...", "info", logger)
            cell_start = ws2.Cells(11, 2)
            cell_end = ws2.Cells(200, 6)
            ws2.Range(cell_start, cell_end).ClearContents()
            _log("Limpieza de rango completada.", "debug", logger)

            _log("Preparando filas de modelos a escribir...", "info", logger)
            rows_to_write = []
            for sku_block in data.get("skus", []):
                sku = sku_block.get("sku", "")
                marca_raw = sku_block.get("marca", "")
                marcas = split_marcas(marca_raw)
                specs = sku_block.get("specs", "")
                modelos = sku_block.get("modelos", [])
                descripcion = data.get("producto", "")

                if not modelos:
                    modelos = [sku]

                _log(f"  - Procesando SKU '{sku}' con marcas {marcas} y {len(modelos)} modelos.", "debug", logger)
                for modelo in modelos:
                    for marca in marcas:
                        rows_to_write.append({
                            "descripcion": descripcion,
                            "marca": marca,
                            "modelo": modelo,
                            "specs": specs
                        })

            _log(f"Total de filas individuales resultantes a escribir: {len(rows_to_write)}", "debug", logger)

            # Obtener marcas únicas en el orden original para agrupar
            unique_brands = []
            for r in rows_to_write:
                b = r["marca"]
                if b not in unique_brands:
                    unique_brands.append(b)

            _log(f"Marcas únicas detectadas en orden: {unique_brands}. Escribiendo marcas agrupadas...", "info", logger)

            # Escribir filas agrupando por marca
            current_row = 11
            for brand in unique_brands:
                _log(f"  -> Escribiendo modelos para marca: '{brand}'", "info", logger)
                for r in rows_to_write:
                    if r["marca"] == brand:
                        _log(f"    * Fila {current_row}: '{r['descripcion']}' | Brand='{r['marca']}' | Model='{r['modelo']}' | Specs='{r['specs']}'", "debug", logger)
                        ws2.Cells(current_row, 2).Value = r["descripcion"]  # Col B (Denominación)
                        ws2.Cells(current_row, 3).Value = r["marca"]        # Col C (Marca)
                        ws2.Cells(current_row, 4).Value = r["modelo"]       # Col D (Modelo)
                        ws2.Cells(current_row, 5).Value = r["specs"]        # Col E (Características técnicas)
                        
                        # Formato Calibri 12
                        cell_range = ws2.Range(ws2.Cells(current_row, 2), ws2.Cells(current_row, 5))
                        cell_range.Font.Name = "Calibri"
                        cell_range.Font.Size = 12
                        
                        current_row += 1
            _log(f"Escritura de modelos finalizada en fila {current_row - 1}.", "info", logger)
        else:
            _log("ATENCIÓN: Hoja de 'detalle' no encontrada en la plantilla. No se escribieron los modelos.", "warning", logger)

        _log("Guardando cambios en el libro como Macro-Enabled (FileFormat=52)...", "debug", logger)
        wb.SaveAs(abs_out, FileFormat=52)
        _log("Cerrando el libro de Excel...", "debug", logger)
        wb.Close(SaveChanges=False)
        _log("Archivo Excel de Lenor guardado exitosamente mediante win32com.", "info", logger)
        return True
    except Exception as e:
        _log(f"Error crítico en automatización win32com para Lenor Excel: {e}", "error", logger)
        if wb:
            try:
                _log("Intentando cerrar el libro sin guardar debido a un error...", "debug", logger)
                wb.Close(SaveChanges=False)
            except Exception:
                pass
        return False
    finally:
        if excel:
            try:
                _log("Cerrando la aplicación de Excel...", "debug", logger)
                excel.Quit()
            except Exception:
                pass
        _log("Liberando recursos de pythoncom (pythoncom.CoUninitialize)...", "debug", logger)
        pythoncom.CoUninitialize()


def generate_lenor_excel_openpyxl(data: dict, out_path: Path, logger: Any = None) -> None:
    """Fallback openpyxl si win32com falla o no está disponible."""
    _log(f"Ejecutando fallback de openpyxl para Lenor Excel. Destino: {out_path.name}", "info", logger)
    try:
        _log(f"Cargando plantilla base: {TEMPLATE_LENOR_XLSX.name} (keep_vba=True)...", "debug", logger)
        wb = openpyxl.load_workbook(str(TEMPLATE_LENOR_XLSX), keep_vba=True)
    except Exception as e:
        _log(f"Error al abrir la plantilla con openpyxl: {e}", "error", logger)
        raise

    ws1_name = next((s for s in wb.sheetnames if "datos" in s.lower()), wb.sheetnames[0])
    ws1 = wb[ws1_name]
    _log(f"Hoja de datos seleccionada: '{ws1.title}'", "debug", logger)

    ws1["G6"] = _today_ddmmyyyy()

    fabrica = data.get("fabrica", "") or "----"
    direccion = data.get("direccion", "") or "----"
    contacto = data.get("contacto", "") or "----"
    email = data.get("email", "") or ""
    telefono = data.get("telefono", "") or ""

    ws1["C40"] = fabrica
    ws1["C42"] = direccion  # Domicilio legal
    ws1["C44"] = direccion  # Domicilio físico
    ws1["C46"] = contacto
    ws1["E46"] = email
    ws1["G46"] = telefono

    _log("Escribiendo metadatos principales en hoja de datos (openpyxl)...", "debug", logger)

    normas_list = _split_normas(data.get("normas", ""))
    reglamento_val = data.get("reglamento", "") or ""
    norma_val = normas_list[0] if normas_list else ""
    ws1["C57"] = reglamento_val
    ws1["C59"] = norma_val

    ws2_name = next((s for s in wb.sheetnames if "detalle" in s.lower()), None)
    if ws2_name:
        ws2 = wb[ws2_name]
        _log(f"Hoja de detalle seleccionada: '{ws2.title}'. Limpiando filas 11 a 200...", "debug", logger)
        for row in range(11, min(200, ws2.max_row + 1)):
            for col in range(2, 7):
                ws2.cell(row, col).value = None

        current_row = 11
        rows_to_write = []
        for sku_block in data.get("skus", []):
            sku = sku_block.get("sku", "")
            marca_raw = sku_block.get("marca", "")
            marcas = split_marcas(marca_raw)
            specs = sku_block.get("specs", "")
            modelos = sku_block.get("modelos", [])
            descripcion = data.get("producto", "")

            if not modelos:
                modelos = [sku]

            for modelo in modelos:
                for marca in marcas:
                    rows_to_write.append({
                        "descripcion": descripcion,
                        "marca": marca,
                        "modelo": modelo,
                        "specs": specs
                    })

        unique_brands = []
        for r in rows_to_write:
            b = r["marca"]
            if b not in unique_brands:
                unique_brands.append(b)

        _log(f"Escribiendo {len(rows_to_write)} filas agrupadas por marca (openpyxl)...", "info", logger)
        for brand in unique_brands:
            _log(f"  -> Escribiendo marca: '{brand}'", "debug", logger)
            for r in rows_to_write:
                if r["marca"] == brand:
                    ws2.cell(current_row, 2).value = r["descripcion"]
                    ws2.cell(current_row, 3).value = r["marca"]
                    ws2.cell(current_row, 4).value = r["modelo"]
                    ws2.cell(current_row, 5).value = r["specs"]
                    
                    # Formato Calibri 12
                    from openpyxl.styles import Font
                    calibri_12 = Font(name="Calibri", size=12)
                    for col in range(2, 6):
                        ws2.cell(current_row, col).font = calibri_12
                        
                    current_row += 1

    _log(f"Guardando archivo Excel destino con openpyxl: {out_path}...", "debug", logger)
    wb.save(str(out_path))
    _log("Archivo Excel de Lenor guardado por openpyxl (fallback).", "info", logger)


def generate_lenor_excel(data: dict, out_path: Path, logger: Any = None) -> None:
    """Orquesta la generación del Excel de Lenor intentando win32com primero."""
    if WIN32_AVAILABLE:
        _log("Intentando generación de Excel de Lenor con automatización win32com...", "info", logger)
        success = generate_lenor_excel_win32(data, out_path, logger)
        if not success:
            _log("La generación win32com de Lenor falló. Elevando excepción estricta para evitar fallback corruptor.", "error", logger)
            raise RuntimeError(
                "Error al automatizar Excel nativo mediante win32com para Lenor. "
                "Por favor, cierre todas las ventanas de Excel abiertas (especialmente si tienen carteles de alerta o diálogos activos) y vuelva a intentarlo."
            )
    else:
        _log("win32com no está instalado en este sistema. Procediendo con openpyxl.", "warning", logger)
        generate_lenor_excel_openpyxl(data, out_path, logger)


# ═════════════════════════════════════════════════════════════════════════════
# 4. GENERADOR QETKRA EXCEL (win32com + fallback openpyxl)
# ═════════════════════════════════════════════════════════════════════════════

def extract_country(direccion_str: str) -> str:
    """
    Extrae el país del fabricante a partir de la dirección.
    Si no se encuentra una palabra clave conocida, devuelve el último fragmento separado por comas
    limpiando códigos postales, o 'CHINA' por defecto.
    """
    if not direccion_str:
        return "CHINA"
    dir_upper = direccion_str.upper()
    if "CHINA" in dir_upper or "P.R.C" in dir_upper or "PRC" in dir_upper or "PEOPLE'S REPUBLIC" in dir_upper:
        return "CHINA"
    if "ESPAÑA" in dir_upper or "SPAIN" in dir_upper:
        return "ESPAÑA"
    if "ITALIA" in dir_upper or "ITALY" in dir_upper:
        return "ITALIA"
    if "ALEMANIA" in dir_upper or "GERMANY" in dir_upper or "DEUTSCHLAND" in dir_upper:
        return "ALEMANIA"
    if "EEUU" in dir_upper or "USA" in dir_upper or "UNITED STATES" in dir_upper:
        return "EEUU"
    if "TAIWAN" in dir_upper:
        return "TAIWAN"
    
    parts = [p.strip() for p in direccion_str.split(",") if p.strip()]
    if parts:
        last_part = parts[-1].upper()
        # Eliminar números que parezcan códigos postales
        last_part = re.sub(r'\b\d{5,}\b', '', last_part).strip()
        if last_part:
            return last_part
    return "CHINA"


def generate_qetkra_excel_win32(data: dict, esquema: str, out_path: Path, logger: Any = None) -> bool:
    """Escribe la solicitud qetkra usando win32com."""
    if not WIN32_AVAILABLE:
        _log("win32com no está instalado/disponible. Saltando a openpyxl.", "warning", logger)
        return False

    _log(f"Abriendo Excel (win32com) para qetkra. Destino: {out_path.name}", "info", logger)
    _log("Inicializando comunicación COM (pythoncom.CoInitialize)...", "debug", logger)
    pythoncom.CoInitialize()
    excel = None
    wb = None
    try:
        _log("Lanzando instancia de Excel en segundo plano (win32com.client.DispatchEx)...", "debug", logger)
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        _log("Instancia de Excel iniciada (Visible=False, DisplayAlerts=False).", "debug", logger)

        # Copiar plantilla al destino
        _log(f"Copiando plantilla {TEMPLATE_qetkra_XLSX.name} a {out_path.name}...", "info", logger)
        shutil.copy2(str(TEMPLATE_qetkra_XLSX), str(out_path))

        abs_out = str(out_path.resolve())
        _log(f"Abriendo archivo Excel destino: {abs_out}...", "debug", logger)
        wb = excel.Workbooks.Open(abs_out)

        # Hoja Solicitud
        _log("Buscando hoja 'solicitud' en el libro...", "debug", logger)
        ws1 = None
        for s in wb.Worksheets:
            if "solicitud" in s.Name.lower():
                ws1 = s
                break
        if not ws1:
            _log("Hoja con 'solicitud' no encontrada, utilizando primera hoja por defecto.", "warning", logger)
            ws1 = wb.Worksheets(1)
        _log(f"Hoja de solicitud seleccionada: '{ws1.Name}'", "debug", logger)

        _log("Escribiendo metadatos principales en hoja de solicitud...", "info", logger)
        ws1.Range("B1").Value = "'" + _today_ddmmyyyy()
        _log(f"  - Fecha hoy en B1: '{_today_ddmmyyyy()}'", "debug", logger)

        fabrica = data.get("fabrica", "") or "----"
        direccion = data.get("direccion", "") or "----"
        contacto = data.get("contacto", "") or ""
        telefono = data.get("telefono", "") or ""
        email = data.get("email", "") or ""
        pais = extract_country(direccion)

        ws1.Range("B16").Value = fabrica
        ws1.Range("B17").Value = direccion
        ws1.Range("B18").Value = pais
        ws1.Range("B19").Value = contacto
        ws1.Range("B20").Value = telefono
        ws1.Range("B21").Value = email
        ws1.Range("B24").Value = fabrica
        ws1.Range("B25").Value = direccion

        _log(f"  - Fábrica (B16, B24): '{fabrica}'", "debug", logger)
        _log(f"  - Dirección (B17, B25): '{direccion}'", "debug", logger)
        _log(f"  - País (B18): '{pais}'", "debug", logger)
        _log(f"  - Contacto (B19): '{contacto}', Tel (B20): '{telefono}', Email (B21): '{email}'", "debug", logger)

        skus = data.get("skus", [])
        primera_marca = skus[0].get("marca", "") if skus else ""
        producto_val = data.get("producto", "")
        ws1.Range("B28").Value = producto_val
        ws1.Range("B29").Value = primera_marca
        ws1.Range("B30").Value = esquema
        ws1.Range("B31").Value = "Tipo (Sistema 2 ISO/IEC 17067)"

        _log(f"  - Producto (B28): '{producto_val}'", "debug", logger)
        _log(f"  - Primera marca (B29): '{primera_marca}'", "debug", logger)
        _log(f"  - Esquema (B30): '{esquema}'", "debug", logger)

        normas_list = _split_normas(data.get("normas", ""))
        ws1.Range("B32").Value = normas_list[0] if len(normas_list) > 0 else ""
        ws1.Range("B33").Value = normas_list[1] if len(normas_list) > 1 else ""
        if len(normas_list) > 3:
            ws1.Range("B34").Value = ", ".join(normas_list[2:])
        elif len(normas_list) == 3:
            ws1.Range("B34").Value = normas_list[2]
        else:
            ws1.Range("B34").Value = ""

        _log(f"  - Normas escritas: {normas_list[:3]}", "debug", logger)

        lab_val = data.get("oec", "") or data.get("laboratorio", "") or "----"
        ws1.Range("B35").Value = lab_val
        _log(f"  - Laboratorio (B35): '{lab_val}'", "debug", logger)

        # Hoja Anexo de Modelos
        _log("Buscando hoja 'anexo' en el libro...", "debug", logger)
        ws2 = None
        for s in wb.Worksheets:
            if "anexo" in s.Name.lower():
                ws2 = s
                break

        if ws2:
            _log(f"Hoja de anexo seleccionada: '{ws2.Name}'", "debug", logger)
            _log("Limpiando celdas del rango B2:N150...", "info", logger)
            cell_start = ws2.Cells(2, 2)
            cell_end = ws2.Cells(150, 14)
            ws2.Range(cell_start, cell_end).ClearContents()
            _log("Limpieza de rango completada.", "debug", logger)

            _log("Escribiendo modelos en la hoja Anexo...", "info", logger)
            current_row = 2
            for sku_block in skus:
                modelos = sku_block.get("modelos", [])
                marca = sku_block.get("marca", "")  # qetkra no duplica marca
                tension = sku_block.get("tension", "") or "---"
                frecuencia = sku_block.get("frecuencia", "") or "---"
                corriente = sku_block.get("corriente", "") or "---"
                potencia = sku_block.get("potencia", "") or "---"
                aislacion = sku_block.get("aislacion", "") or "---"
                tension_salida = sku_block.get("tension_salida", "") or "---"
                corriente_salida = sku_block.get("corriente_salida", "") or "---"
                grado_ip = sku_block.get("grado_ip", "") or "---"
                casquillo = sku_block.get("casquillo", "") or "---"
                adicional = sku_block.get("adicional", "") or "---"

                if not modelos:
                    modelos = [sku_block.get("sku", "")]

                _log(f"  - Procesando SKU '{sku_block.get('sku')}' con marcas {marca} y {len(modelos)} modelos.", "debug", logger)
                for i, modelo in enumerate(modelos):
                    is_principal = (i == 0)
                    _log(f"    * Fila {current_row}: Principal={modelo if is_principal else '----'} | Alternativo={'----' if is_principal else modelo} | Brand='{marca}'", "debug", logger)
                    ws2.Cells(current_row, 2).Value = modelo if is_principal else "----"  # Col B
                    ws2.Cells(current_row, 3).Value = "----" if is_principal else modelo  # Col C
                    ws2.Cells(current_row, 4).Value = marca                              # Col D (Marca)
                    ws2.Cells(current_row, 5).Value = tension                            # Col E
                    ws2.Cells(current_row, 6).Value = frecuencia                         # Col F
                    ws2.Cells(current_row, 7).Value = corriente                          # Col G
                    ws2.Cells(current_row, 8).Value = potencia                           # Col H
                    ws2.Cells(current_row, 9).Value = aislacion                          # Col I
                    ws2.Cells(current_row, 10).Value = tension_salida                    # Col J
                    ws2.Cells(current_row, 11).Value = corriente_salida                  # Col K
                    ws2.Cells(current_row, 12).Value = grado_ip                          # Col L
                    ws2.Cells(current_row, 13).Value = casquillo                         # Col M
                    ws2.Cells(current_row, 14).Value = adicional                         # Col N
                    current_row += 1
            _log(f"Escritura de modelos finalizada en fila {current_row - 1}.", "info", logger)
        else:
            _log("ATENCIÓN: Hoja de 'anexo' no encontrada en la plantilla. No se escribieron los modelos.", "warning", logger)

        _log("Guardando cambios en el libro...", "debug", logger)
        wb.Save()
        _log("Cerrando el libro de Excel...", "debug", logger)
        wb.Close(SaveChanges=True)
        _log("Archivo Excel de qetkra guardado exitosamente mediante win32com.", "info", logger)
        return True
    except Exception as e:
        _log(f"Error crítico en automatización win32com para qetkra Excel: {e}", "error", logger)
        if wb:
            try:
                _log("Intentando cerrar el libro sin guardar debido a un error...", "debug", logger)
                wb.Close(SaveChanges=False)
            except Exception:
                pass
        return False
    finally:
        if excel:
            try:
                _log("Cerrando la aplicación de Excel...", "debug", logger)
                excel.Quit()
            except Exception:
                pass
        _log("Liberando recursos de pythoncom (pythoncom.CoUninitialize)...", "debug", logger)
        pythoncom.CoUninitialize()


def generate_qetkra_excel_openpyxl(data: dict, esquema: str, out_path: Path, logger: Any = None) -> None:
    """Fallback openpyxl si win32com falla o no está disponible."""
    _log(f"Ejecutando fallback de openpyxl para qetkra Excel. Destino: {out_path.name}", "info", logger)
    try:
        _log(f"Cargando plantilla base: {TEMPLATE_qetkra_XLSX.name}...", "debug", logger)
        wb = openpyxl.load_workbook(str(TEMPLATE_qetkra_XLSX))
    except Exception as e:
        _log(f"Error al abrir la plantilla con openpyxl: {e}", "error", logger)
        raise

    ws_sol_name = next((s for s in wb.sheetnames if "solicitud" in s.lower()), wb.sheetnames[0])
    ws = wb[ws_sol_name]
    _log(f"Hoja de solicitud seleccionada: '{ws.title}'", "debug", logger)
    ws["B1"] = _today_ddmmyyyy()

    fabrica = data.get("fabrica", "") or "----"
    direccion = data.get("direccion", "") or "----"
    contacto = data.get("contacto", "") or ""
    telefono = data.get("telefono", "") or ""
    email = data.get("email", "") or ""
    pais = extract_country(direccion)

    ws["B16"] = fabrica
    ws["B17"] = direccion
    ws["B18"] = pais
    ws["B19"] = contacto
    ws["B20"] = telefono
    ws["B21"] = email
    ws["B24"] = fabrica
    ws["B25"] = direccion

    skus = data.get("skus", [])
    primera_marca = skus[0].get("marca", "") if skus else ""
    ws["B28"] = data.get("producto", "")
    ws["B29"] = primera_marca
    ws["B30"] = esquema
    ws["B31"] = "Tipo (Sistema 2 ISO/IEC 17067)"

    _log("Escribiendo metadatos principales en hoja de solicitud (openpyxl)...", "debug", logger)

    normas_list = _split_normas(data.get("normas", ""))
    ws["B32"] = normas_list[0] if len(normas_list) > 0 else ""
    ws["B33"] = normas_list[1] if len(normas_list) > 1 else ""
    if len(normas_list) > 3:
        ws["B34"] = ", ".join(normas_list[2:])
    elif len(normas_list) == 3:
        ws["B34"] = normas_list[2]
    else:
        ws["B34"] = ""

    ws["B35"] = data.get("oec", "") or data.get("laboratorio", "") or "----"

    ws_anx_name = next((s for s in wb.sheetnames if "anexo" in s.lower()), None)
    if ws_anx_name:
        ws_anx = wb[ws_anx_name]
        _log(f"Hoja de anexo seleccionada: '{ws_anx.title}'. Limpiando filas 2 a 150 (columnas B a N)...", "debug", logger)
        for row in range(2, min(150, ws_anx.max_row + 1)):
            for col in range(2, 15):
                ws_anx.cell(row, col).value = None

        current_row = 2
        for sku_block in skus:
            modelos = sku_block.get("modelos", [])
            marca = sku_block.get("marca", "")
            tension = sku_block.get("tension", "") or "---"
            frecuencia = sku_block.get("frecuencia", "") or "---"
            corriente = sku_block.get("corriente", "") or "---"
            potencia = sku_block.get("potencia", "") or "---"
            aislacion = sku_block.get("aislacion", "") or "---"
            tension_salida = sku_block.get("tension_salida", "") or "---"
            corriente_salida = sku_block.get("corriente_salida", "") or "---"
            grado_ip = sku_block.get("grado_ip", "") or "---"
            casquillo = sku_block.get("casquillo", "") or "---"
            adicional = sku_block.get("adicional", "") or "---"

            if not modelos:
                modelos = [sku_block.get("sku", "")]

            _log(f"  - Escribiendo {len(modelos)} modelos para SKU '{sku_block.get('sku')}' (openpyxl)...", "debug", logger)
            for i, modelo in enumerate(modelos):
                is_principal = (i == 0)
                ws_anx.cell(current_row, 2).value = modelo if is_principal else "----"
                ws_anx.cell(current_row, 3).value = "----" if is_principal else modelo
                ws_anx.cell(current_row, 4).value = marca
                ws_anx.cell(current_row, 5).value = tension
                ws_anx.cell(current_row, 6).value = frecuencia
                ws_anx.cell(current_row, 7).value = corriente
                ws_anx.cell(current_row, 8).value = potencia
                ws_anx.cell(current_row, 9).value = aislacion
                ws_anx.cell(current_row, 10).value = tension_salida
                ws_anx.cell(current_row, 11).value = corriente_salida
                ws_anx.cell(current_row, 12).value = grado_ip
                ws_anx.cell(current_row, 13).value = casquillo
                ws_anx.cell(current_row, 14).value = adicional
                current_row += 1

    _log(f"Guardando archivo Excel destino con openpyxl: {out_path}...", "debug", logger)
    wb.save(str(out_path))
    _log("Archivo Excel de qetkra guardado por openpyxl (fallback).", "info", logger)


def generate_qetkra_excel(data: dict, esquema: str, out_path: Path, logger: Any = None) -> None:
    """Orquesta la generación del Excel de qetkra."""
    if WIN32_AVAILABLE:
        _log("Intentando generación de Excel de qetkra con automatización win32com...", "info", logger)
        success = generate_qetkra_excel_win32(data, esquema, out_path, logger)
        if not success:
            _log("La generación win32com de qetkra falló. Elevando excepción estricta para evitar fallback corruptor.", "error", logger)
            raise RuntimeError(
                "Error al automatizar Excel nativo mediante win32com para qetkra. "
                "Por favor, cierre todas las ventanas de Excel abiertas (especialmente si tienen carteles de alerta o diálogos activos) y vuelva a intentarlo."
            )
    else:
        _log("win32com no está instalado en este sistema. Procediendo con openpyxl.", "warning", logger)
        generate_qetkra_excel_openpyxl(data, esquema, out_path, logger)


# ═════════════════════════════════════════════════════════════════════════════
# 5. REEMPLAZO DE TAGS EN WORD (python-docx respetando runs)
# ═════════════════════════════════════════════════════════════════════════════

def _set_cell_text(cell, text: str, font_name: str | None = None, font_size_pt: int | float | None = None, bold: bool | None = None) -> None:
    """Escribe texto en una celda de tabla Word preservando el formato del primer run si existe o aplicando uno nuevo."""
    from docx.shared import Pt
    run = None
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
            run = p.runs[0]
        else:
            run = p.add_run(text)
        for extra_p in cell.paragraphs[1:]:
            p_elem = extra_p._p
            p_elem.getparent().remove(p_elem)
    else:
        p = cell.add_paragraph()
        run = p.add_run(text)
        
    if run:
        if font_name:
            run.font.name = font_name
        if font_size_pt:
            run.font.size = Pt(font_size_pt)
        if bold is not None:
            run.font.bold = bold


def replace_tags_in_docx(doc: Document, replacements: dict[str, str]) -> None:
    """Reemplaza etiquetas fijas ({FECHA}, {PRODUCTO}) en párrafos y tablas preservando estilos."""
    for paragraph in doc.paragraphs:
        for tag, val in replacements.items():
            if tag in paragraph.text:
                replaced = False
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, val)
                        replaced = True
                if not replaced and paragraph.runs:
                    full_text = "".join(r.text for r in paragraph.runs)
                    new_text = full_text.replace(tag, val)
                    paragraph.runs[0].text = new_text
                    for r in paragraph.runs[1:]:
                        r.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for tag, val in replacements.items():
                        if tag in paragraph.text:
                            replaced = False
                            for run in paragraph.runs:
                                if tag in run.text:
                                    run.text = run.text.replace(tag, val)
                                    replaced = True
                            if not replaced and paragraph.runs:
                                full_text = "".join(r.text for r in paragraph.runs)
                                new_text = full_text.replace(tag, val)
                                paragraph.runs[0].text = new_text
                                for r in paragraph.runs[1:]:
                                    r.text = ""


# ═════════════════════════════════════════════════════════════════════════════
# 6. GENERADORES WORD
# ═════════════════════════════════════════════════════════════════════════════

def generate_lenor_word(data: dict, out_path: Path, logger: Any = None) -> None:
    """Genera la nota comercial de Lenor a partir del Word base con tags."""
    _log(f"Iniciando generación de Nota Comercial Word para Lenor. Destino: {out_path.name}", "info", logger)
    try:
        _log(f"Abriendo plantilla Word: {TEMPLATE_LENOR_DOCX.name}...", "debug", logger)
        doc = Document(str(TEMPLATE_LENOR_DOCX))
    except Exception as e:
        _log(f"Error al cargar la plantilla Word: {e}", "error", logger)
        raise

    fecha_hoy = _today_long_es()
    _log(f"Reemplazando etiqueta {{FECHA}} por '{fecha_hoy}'...", "debug", logger)
    replace_tags_in_docx(doc, {"{FECHA}": fecha_hoy})

    if not doc.tables:
        _log("ADVERTENCIA: No se encontraron tablas en la plantilla Word de Lenor. Guardando archivo tal cual.", "warning", logger)
        _log(f"Guardando archivo Word: {out_path}...", "debug", logger)
        doc.save(str(out_path))
        return

    table = doc.tables[0]
    _log("Tabla de modelos detectada en Word. Limpiando filas previas de ejemplo (manteniendo cabecera)...", "debug", logger)
    ref_row = table.rows[1] if len(table.rows) > 1 else None

    # Limpiar filas previas
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    _log("Filas previas limpiadas correctamente.", "debug", logger)

    # Construir filas sin duplicar marcas (solo Lenor)
    rows_data = []
    for sku_block in data.get("skus", []):
        sku = sku_block.get("sku", "")
        marca_raw = sku_block.get("marca", "")
        specs = sku_block.get("specs", "")
        modelos = sku_block.get("modelos", [])

        if not modelos:
            modelos = [sku]

        for modelo in modelos:
            rows_data.append((sku, marca_raw, modelo, specs))

    _log(f"Escribiendo {len(rows_data)} filas de modelos en la tabla del Word de Lenor...", "info", logger)
    # Escribir filas
    for sku_val, marca_val, modelo_val, specs_val in rows_data:
        _log(f"  * Escribiendo fila Word: SKU='{sku_val}' | Marcas='{marca_val}' | Modelo='{modelo_val}'", "debug", logger)
        if ref_row:
            new_tr = deepcopy(ref_row._tr)
            table._tbl.append(new_tr)
            new_row = table.rows[-1]
            cells = new_row.cells
        else:
            new_row = table.add_row()
            cells = new_row.cells

        _set_cell_text(cells[0], sku_val, font_name="Calibri", font_size_pt=12)
        _set_cell_text(cells[1], marca_val, font_name="Calibri", font_size_pt=12)
        _set_cell_text(cells[2], modelo_val, font_name="Calibri", font_size_pt=12)
        if len(cells) > 3:
            _set_cell_text(cells[3], specs_val, font_name="Calibri", font_size_pt=12)

    _log(f"Guardando Nota comercial Word generada en: {out_path}...", "debug", logger)
    doc.save(str(out_path))
    _log("Nota comercial de Lenor Word generada exitosamente.", "info", logger)


def generate_qetkra_word(data: dict, out_path: Path, svg_bytes: bytes | None = None, logger: Any = None) -> None:
    """Genera la nota de aclaración de qetkra a partir del Word base con tags."""
    _log(f"Iniciando generación de Nota de Aclaración Word para qetkra. Destino: {out_path.name}", "info", logger)
    try:
        _log(f"Abriendo plantilla Word: {TEMPLATE_qetkra_DOCX.name}...", "debug", logger)
        doc = Document(str(TEMPLATE_qetkra_DOCX))
    except Exception as e:
        _log(f"Error al cargar la plantilla Word: {e}", "error", logger)
        raise

    # Si hay svg_bytes, convertimos a PNG y reemplazamos el QR en el Word
    if svg_bytes:
        _log("SVG de código QR provisto. Intentando reemplazar en Nota Word de qetkra...", "info", logger)
        try:
            import fitz
            svg_text = svg_bytes.decode("utf-8", errors="ignore")

            # Extraer viewBox para establecer width y height absolutos
            match = re.search(r'viewBox="\s*0\s+0\s+([\d.]+)\s+([\d.]+)\s*"', svg_text)
            if match:
                vb_w = float(match.group(1))
                vb_h = float(match.group(2))
                svg_text = re.sub(r'width="[^"]+"', f'width="{vb_w}"', svg_text, count=1)
                svg_text = re.sub(r'height="[^"]+"', f'height="{vb_h}"', svg_text, count=1)

            # Escalar el QR para que ocupe todo el cuadrado (quitando traslación y aumentando la escala de 1.3 a 1.9462)
            svg_text = re.sub(
                r'transform="translate\(\s*\d+\s*,\s*\d+\s*\)\s+scale\(\s*[\d.]+\s*\)"',
                'transform="scale(1.9462)"',
                svg_text
            )
            svg_text = svg_text.replace('transform="translate(7, 7) scale(1.3)"', 'transform="scale(1.9462)"')
            svg_bytes_scaled = svg_text.encode("utf-8")

            svg_doc = fitz.open(stream=svg_bytes_scaled, filetype="svg")
            page = svg_doc[0]
            pix = page.get_pixmap(dpi=300)
            png_bytes = pix.tobytes("png")
            
            # Reemplazar la imagen en la plantilla
            replaced = False
            # Primero intentar en Paragraph 8 si contiene un blip
            if len(doc.paragraphs) > 8:
                p8 = doc.paragraphs[8]
                blips = p8._element.xpath('.//a:blip')
                if blips:
                    embed_id = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id in doc.part.rels:
                        rel = doc.part.rels[embed_id]
                        if "image1.png" in rel.target_ref or len(rel.target_part.blob) == 8761:
                            rel.target_part._blob = png_bytes
                            replaced = True
                            _log("QR reemplazado usando Paragraph 8 en Nota Word.", "info", logger)
            
            # Fallback o búsqueda directa por relación si no se reemplazó
            if not replaced:
                for rel in doc.part.rels.values():
                    if "image1.png" in rel.target_ref or len(rel.target_part.blob) == 8761:
                        rel.target_part._blob = png_bytes
                        replaced = True
                        _log("QR reemplazado usando relación directa (image1.png / 8761 bytes) en Nota Word.", "info", logger)
                        break
                        
            if not replaced:
                _log("No se pudo encontrar el marcador de posición del QR en el Word para reemplazarlo.", "warning", logger)
        except Exception as ex:
            _log(f"Error al renderizar o reemplazar el QR en el Word de qetkra: {ex}", "error", logger)

    fecha_hoy = _today_long_es()
    producto = data.get("producto", "")
    
    _log(f"Reemplazando etiquetas {{FECHA}} por '{fecha_hoy}' y {{PRODUCTO}} por '{producto}'...", "debug", logger)
    replace_tags_in_docx(doc, {
        "{FECHA}": fecha_hoy,
        "{PRODUCTO}": producto
    })

    # Determinar clase de aislación para la ficha
    skus = data.get("skus", [])
    insulation_class = ""
    if skus:
        insulation_class = skus[0].get("aislacion", "") or ""
    
    insulation_upper = insulation_class.upper()
    is_class_ii = False
    is_class_i = False
    
    if "CLASE II" in insulation_upper or "CLASS II" in insulation_upper or "CLASE 2" in insulation_upper or "CLASS 2" in insulation_upper:
        is_class_ii = True
    elif "CLASE I" in insulation_upper or "CLASS I" in insulation_upper or "CLASE 1" in insulation_upper or "CLASS 1" in insulation_upper:
        is_class_i = True

    ficha_text = "El producto incluirá una ficha de alimentación Formato IRAM 2063 o IRAM 2073 certificada, según corresponda a su clase de aislación."
    if is_class_i:
        ficha_text = "El producto incluirá una ficha de alimentación Formato IRAM 2073 certificada, correspondiente a su clase de aislación."
        _log("Clase de aislación detectada como Clase I. Reemplazando texto de ficha por IRAM 2073.", "info", logger)
    elif is_class_ii:
        ficha_text = "El producto incluirá una ficha de alimentación Formato IRAM 2063 certificada, correspondiente a su clase de aislación."
        _log("Clase de aislación detectada como Clase II. Reemplazando texto de ficha por IRAM 2063.", "info", logger)
    else:
        _log("No se pudo detectar clase de aislación clara (o no es I ni II). Manteniendo texto de ficha por defecto.", "info", logger)

    target_phrase = "El producto incluirá una ficha de alimentación Formato IRAM 2063 o IRAM 2073"
    for paragraph in doc.paragraphs:
        if target_phrase in paragraph.text:
            if paragraph.runs:
                paragraph.runs[0].text = ficha_text
                for r in paragraph.runs[1:]:
                    r.text = ""
            else:
                paragraph.text = ficha_text
            break

    if not doc.tables:
        _log("ADVERTENCIA: No se encontraron tablas en la plantilla Word de qetkra. Guardando archivo tal cual.", "warning", logger)
        _log(f"Guardando archivo Word: {out_path}...", "debug", logger)
        doc.save(str(out_path))
        return

    table = doc.tables[0]
    _log("Tabla de correlatividad detectada en Word. Limpiando filas previas de ejemplo...", "debug", logger)
    ref_row = table.rows[1] if len(table.rows) > 1 else None

    # Limpiar filas previas
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    _log("Filas previas limpiadas correctamente.", "debug", logger)

    # Construir filas (Modelo BIDCOM | Modelo CB)
    # qetkra no duplica filas por marcas múltiples
    rows_data = []
    for sku_block in data.get("skus", []):
        modelos = sku_block.get("modelos", [])
        modelo_fab = sku_block.get("modelo_fabrica", "") or "---"

        if not modelos:
            modelos = [sku_block.get("sku", "")]

        for modelo in modelos:
            rows_data.append((modelo, modelo_fab))

    _log(f"Escribiendo {len(rows_data)} filas en la tabla de correlatividad de qetkra...", "info", logger)
    for modelo_val, fab_val in rows_data:
        _log(f"  * Escribiendo fila Word: Bidcom='{modelo_val}' | Fábrica/CB='{fab_val}'", "debug", logger)
        if ref_row:
            new_tr = deepcopy(ref_row._tr)
            table._tbl.append(new_tr)
            new_row = table.rows[-1]
            cells = new_row.cells
        else:
            new_row = table.add_row()
            cells = new_row.cells

        _set_cell_text(cells[0], modelo_val)
        if len(cells) > 1:
            _set_cell_text(cells[1], fab_val)

    _log(f"Guardando Nota Word de qetkra en: {out_path}...", "debug", logger)
    doc.save(str(out_path))
    _log("Nota de aclaración de qetkra Word generada exitosamente.", "info", logger)


def generate_tuv_word(data: dict, out_path: Path, logger: Any = None) -> None:
    """Genera la Solicitud de Servicio oficial de TÜV Rheinland en Word (.docx)."""
    _log(f"Iniciando generación de Solicitud de Servicio Word para TÜV Rheinland. Destino: {out_path.name}", "info", logger)
    try:
        _log(f"Abriendo plantilla Word: {TEMPLATE_TUV_DOCX.name}...", "debug", logger)
        doc = Document(str(TEMPLATE_TUV_DOCX))
    except Exception as e:
        _log(f"Error al cargar la plantilla Word de TÜV: {e}", "error", logger)
        raise

    if not doc.tables:
        _log("ERROR: No se encontró la tabla de solicitud en la plantilla de TÜV.", "error", logger)
        raise ValueError("Plantilla de TÜV inválida: no contiene tabla.")

    table = doc.tables[0]

    # Extraer campos
    producto = data.get("producto", "") or "---"
    motivo = data.get("motivo", "") or "Adición de modelo"
    reglamento = data.get("reglamento", "") or "Resolución S.I.C. N° 16/2025"
    esquema = data.get("esquema", "") or "Tipo 2"
    normas = data.get("normas", "") or "---"
    laboratorio = data.get("laboratorio", "") or "TÜV Rheinland"
    certificado = data.get("certificado", "") or ""
    fabrica_nombre = data.get("fabrica", "") or "---"
    fabrica_direccion = data.get("direccion", "") or "---"

    # Modelos y marcas
    all_models = []
    all_marcas = set()
    for b in data.get("skus", []):
        if b.get("marca"):
            for m in split_marcas(b.get("marca")):
                all_marcas.add(m)
        mods = b.get("modelos", [])
        if mods:
            all_models.extend(mods)
        elif b.get("sku"):
            all_models.append(b.get("sku"))

    seen_m = set()
    unique_models = [m for m in all_models if not (m in seen_m or seen_m.add(m))]
    modelos_str = ", ".join(unique_models) if unique_models else "---"
    marcas_str = ", ".join(sorted(all_marcas)) if all_marcas else "GADNIC"

    _log("Completando datos de Titular, Solicitante y Facturación en la tabla...", "debug", logger)
    # 1. Titular (Fila 1 y 2)
    _set_cell_text(table.rows[1].cells[2], "BIDCOM S.R.L.", font_name="Calibri", font_size_pt=10, bold=True)
    _set_cell_text(table.rows[1].cells[7], "30-71106936-0", font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[2].cells[2], "Bouchard 468, 5° I, CABA. CP 1004", font_name="Calibri", font_size_pt=10)

    # 2. Fábrica (Fila 4 y 5)
    _log(f"Completando datos de Fábrica: {fabrica_nombre}...", "debug", logger)
    _set_cell_text(table.rows[4].cells[2], fabrica_nombre, font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[5].cells[2], fabrica_direccion, font_name="Calibri", font_size_pt=10)

    # 3. Facturación (Fila 7)
    _set_cell_text(table.rows[7].cells[0], "Razón social a facturar:    BIDCOM S.R.L.", font_name="Calibri", font_size_pt=10)

    # 4. Solicitante (Fila 10, 11, 12, 13)
    _set_cell_text(table.rows[10].cells[2], "BIDCOM S.R.L.", font_name="Calibri", font_size_pt=10, bold=True)
    _set_cell_text(table.rows[10].cells[6], "CUIT: 30-71106936-0", font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[11].cells[2], "Bouchard 468, 5° I, CABA. CP 1004", font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[12].cells[0], "Persona de contacto : Federico Dean / Ariana Jallinsky", font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[12].cells[5], "Dept.: COMEX", font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[13].cells[0], "Tel.: 3960-0184", font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[13].cells[2], "Ext.: ---", font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[13].cells[5], "e-mail: federico.de@bidcom.com.ar / ariana.jallinsky@bidcom.com.ar", font_name="Calibri", font_size_pt=9)

    # 5. Datos Técnicos (Fila 15, 16, 17, 18, 19, 20)
    _log(f"Completando datos técnicos: Producto='{producto}', Modelos='{modelos_str}', Marcas='{marcas_str}'...", "debug", logger)
    _set_cell_text(table.rows[15].cells[2], producto, font_name="Calibri", font_size_pt=10, bold=True)
    _set_cell_text(table.rows[16].cells[2], modelos_str, font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[17].cells[2], marcas_str, font_name="Calibri", font_size_pt=10)
    _set_cell_text(table.rows[18].cells[2], normas, font_name="Calibri", font_size_pt=10)

    if certificado and certificado != "SIN_NRO":
        _set_cell_text(table.rows[19].cells[0], f"¿El producto ya fue certificado previamente por TÜV?    No   Si [X]:  Indicar Nº Certificado: {certificado}", font_name="Calibri", font_size_pt=10)
    else:
        _set_cell_text(table.rows[19].cells[0], "¿El producto ya fue certificado previamente por TÜV?   [X] No   Si:  Indicar Nº Certificado: ", font_name="Calibri", font_size_pt=10)

    _set_cell_text(table.rows[20].cells[2], laboratorio or "TÜV Rheinland Argentina S.A.", font_name="Calibri", font_size_pt=10)

    # 6. Servicios Solicitados / Checkboxes en Fila 22
    _log("Marcando servicios y resoluciones solicitadas...", "debug", logger)
    r22 = table.rows[22]
    reg_clean = reglamento.lower()
    esq_clean = esquema.lower()
    mot_clean = motivo.lower()

    for p in r22.cells[0].paragraphs:
        txt = p.text.strip()
        if "16/25" in txt or "16/2025" in txt:
            if "16/25" in reg_clean or "16/2025" in reg_clean or "seguridad" in reg_clean:
                p.text = "[X] Resolución S.I.C. N° 16/25 y mod."
        elif "236/24" in txt or "236/2024" in txt:
            if "236/24" in reg_clean or "236/2024" in reg_clean:
                p.text = "[X] Resolución S.I.C. N° 236/24 y mod."
        elif "17/25" in txt or "17/2025" in txt:
            if "17/25" in reg_clean or "17/2025" in reg_clean:
                p.text = "[X] Resolución S.I.C. N° 17/25 y mod."
        elif "Apéndice" in txt:
            prod_low = producto.lower()
            if ("fuente" in prod_low or "cargador" in prod_low) and "Apéndice I " in txt:
                p.text = "\t [X] Apéndice I (Fuentes y cargadores)"
            elif ("iluminac" in prod_low or "lampar" in prod_low or "led" in prod_low) and "Apéndice III" in txt:
                p.text = "\t [X] Apéndice III (Dispositivos de iluminación)"
            elif ("audio" in prod_low or "video" in prod_low or "tv" in prod_low or "parlante" in prod_low or "auricular" in prod_low) and "Apéndice IV" in txt:
                p.text = "\t [X] Apéndice IV (Aparatos de electrónica, audio y video)"
            elif "Apéndice II" in txt:
                p.text = "\t [X] Apéndice II (Aparatos eléctricos de uso doméstico)"
        elif "Esquema tipo 2" in txt:
            if "2" in esq_clean or "tipo" in esq_clean:
                p.text = "[X] Certificación de tipo (Esquema tipo 2 *)"
        elif "Esquema tipo 5" in txt:
            if "5" in esq_clean or "marca" in esq_clean:
                p.text = "[X] Certificación de marca (Esquema tipo 5 *)"
        elif "Esquema tipo 1b" in txt:
            if "lote" in esq_clean or "1b" in esq_clean:
                p.text = "[X] Certificación de lote (Esquema tipo 1b *)"

    for p in r22.cells[3].paragraphs:
        txt = p.text.strip()
        if "Adición de modelo" in txt and ("adicion" in mot_clean or "ampliacion" in mot_clean or "modelo" in mot_clean or "nuevo" in mot_clean):
            p.text = "[X] Adición de modelo"
        elif "Actualización de norma" in txt and "norma" in mot_clean:
            p.text = "[X] Actualización de norma"
        elif "Cambio de datos" in txt and "cambio" in mot_clean:
            p.text = "[X] Cambio de datos técnicos"

    # Fecha en párrafos
    fecha_hoy = date.today().strftime("%d/%m/%Y")
    for p in doc.paragraphs:
        if "Fecha:" in p.text:
            p.text = f"Fecha: {fecha_hoy}"

    _log(f"Guardando Solicitud Word de TÜV en: {out_path}...", "debug", logger)
    doc.save(str(out_path))
    _log("Solicitud de Servicio de TÜV Rheinland generada exitosamente.", "info", logger)


# ═════════════════════════════════════════════════════════════════════════════
# 7. GENERADOR PDF QR (SVG → PDF A4 vectorial)
# ═════════════════════════════════════════════════════════════════════════════

MM_TO_PT = 2.8346
QR_SIZE_SMALL = (12.56, 15)   # mm (mantiene aspecto 72/86)
QR_SIZE_LARGE = (25.12, 30)   # mm (mantiene aspecto 72/86)

A4_W, A4_H = 595, 842

QR_POSITIONS = [
    (QR_SIZE_SMALL[0], QR_SIZE_SMALL[1], 60,  200, f"12.5×15mm"),
    (QR_SIZE_LARGE[0], QR_SIZE_LARGE[1], 200, 200, f"25×30mm"),
]


def generate_qr_pdf(svg_bytes: bytes, nro: str, out_path: Path, logger: Any = None) -> None:
    """Genera un PDF A4 con el QR SVG insertado vectorialmente (sin pixelar)."""
    try:
        import fitz  # PyMuPDF
        _log("Librería PyMuPDF (fitz) importada con éxito.", "debug", logger)
    except ImportError:
        _log("Error crítico: PyMuPDF no está instalado.", "error", logger)
        raise ImportError("PyMuPDF no está instalado. Instalar con: pip install pymupdf")

    _log(f"Generando PDF QR vectorial para Certificado {nro}. Destino: {out_path.name}", "info", logger)
    doc = fitz.open()
    page = doc.new_page(width=A4_W, height=A4_H)
    _log(f"Creada página A4 (ancho={A4_W}pt, alto={A4_H}pt).", "debug", logger)

    page.insert_text(
        (60, 60),
        f"QR Certificado {nro}",
        fontsize=16,
        color=(0.1, 0.1, 0.1),
    )

    # Cargar SVG como string para realizar la escala del QR de forma vectorial
    _log("Decodificando y procesando SVG del QR para adaptar la escala...", "debug", logger)
    svg_text = svg_bytes.decode("utf-8", errors="ignore")

    # Extraer viewBox para establecer width y height absolutos en el tag svg y evitar que PyMuPDF use Letter por defecto
    match = re.search(r'viewBox="\s*0\s+0\s+([\d.]+)\s+([\d.]+)\s*"', svg_text)
    if match:
        vb_w = float(match.group(1))
        vb_h = float(match.group(2))
        _log(f"viewBox detectado en SVG: {vb_w}x{vb_h}. Sobreestableciendo width y height en tag SVG...", "debug", logger)
        svg_text = re.sub(r'width="[^"]+"', f'width="{vb_w}"', svg_text, count=1)
        svg_text = re.sub(r'height="[^"]+"', f'height="{vb_h}"', svg_text, count=1)

    # Si viene con la traslación y escala por defecto, la aumentamos para que
    # llene completamente el rectángulo gris/blanco del QR en la etiqueta.
    # El QR es de 29x29. El viewBox es 56.44. 56.44 / 29 = 1.9462
    svg_text = re.sub(
        r'transform="translate\(\s*\d+\s*,\s*\d+\s*\)\s+scale\(\s*[\d.]+\s*\)"',
        'transform="scale(1.9462)"',
        svg_text
    )
    svg_text = svg_text.replace('transform="translate(7, 7) scale(1.3)"', 'transform="scale(1.9462)"')
    svg_bytes_scaled = svg_text.encode("utf-8")
    _log("Escalado vectorial de SVG QR completado (transform='scale(1.9462)').", "debug", logger)

    # Cargar SVG como doc temporal
    _log("Cargando stream SVG en PyMuPDF...", "debug", logger)
    svg_doc = fitz.open(stream=svg_bytes_scaled, filetype="svg")
    # Convertir SVG a PDF (conserva los vectores puros)
    _log("Convirtiendo SVG a representación PDF vectorial...", "debug", logger)
    svg_pdf_bytes = svg_doc.convert_to_pdf()
    svg_pdf_doc = fitz.open("pdf", svg_pdf_bytes)

    for mm_w, mm_h, x_pt, y_pt, label in QR_POSITIONS:
        w_pt = mm_w * MM_TO_PT
        h_pt = mm_h * MM_TO_PT
        rect = fitz.Rect(x_pt, y_pt, x_pt + w_pt, y_pt + h_pt)
        _log(f"  - Insertando QR {label} en rectángulo (x={x_pt}, y={y_pt}, w={w_pt:.2f}pt, h={h_pt:.2f}pt)...", "debug", logger)

        # Insertar página PDF vectorialmente
        page.show_pdf_page(rect, svg_pdf_doc, 0, keep_proportion=True)

        # Etiqueta
        page.insert_text(
            (x_pt, y_pt + h_pt + 14),
            label,
            fontsize=8,
            color=(0.4, 0.4, 0.4),
        )

    svg_pdf_doc.close()
    svg_doc.close()
    _log("Documentos temporales cerrados.", "debug", logger)
    
    pdf_bytes = doc.tobytes()
    doc.close()

    _log(f"Guardando archivo PDF en disco: {out_path}...", "debug", logger)
    out_path.write_bytes(pdf_bytes)
    _log(f"PDF QR vectorial guardado con éxito.", "info", logger)


def generate_lenor_datasheet_excel(data: dict, out_path: Path, logger: Any = None) -> None:
    """Genera la planilla de fotos (datasheet) precargada para Lenor."""
    _log(f"Generando planilla de fotos (datasheet) para Lenor. Destino: {out_path.name}", "info", logger)
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Fichas técnicas"
        ws.views.sheetView[0].showGridLines = True

        from openpyxl.styles import Font, Border, Side, Alignment
        
        # Borde negro para la tabla
        thin_side = Side(style='thin', color='000000')
        thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
        
        header_font = Font(name="Calibri", size=11, bold=True)
        data_font = Font(name="Calibri", size=10)
        
        current_row = 3
        for sku_block in data.get("skus", []):
            sku = sku_block.get("sku", "")
            marca_raw = sku_block.get("marca", "")
            specs = sku_block.get("specs", "")
            modelos = sku_block.get("modelos", [])
            if not modelos:
                modelos = [sku]
                
            # Cabecera
            headers = ["Marca", "Modelos", "Caracteristicas tecnicas", "Imagen"]
            for col_idx, text in enumerate(headers, 2):
                cell = ws.cell(current_row, col_idx)
                cell.value = text
                cell.font = header_font
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
            current_row += 1
            start_data_row = current_row
            
            # Escribir y dar formato individual a cada celda de datos
            for idx, modelo in enumerate(modelos):
                row_idx = start_data_row + idx
                
                # Valores en celdas individuales (openpyxl los requiere en la celda superior izquierda al combinar)
                if idx == 0:
                    ws.cell(row_idx, 2).value = marca_raw
                    ws.cell(row_idx, 4).value = specs
                else:
                    ws.cell(row_idx, 2).value = ""
                    ws.cell(row_idx, 4).value = ""
                    
                ws.cell(row_idx, 3).value = modelo
                ws.cell(row_idx, 5).value = "" # Columna de Imagen vacía
                
                # Aplicar fuente, bordes y alineaciones individuales
                for col_idx in range(2, 6):
                    cell = ws.cell(row_idx, col_idx)
                    cell.font = data_font
                    cell.border = thin_border
                    if col_idx == 4:
                        cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
                    else:
                        cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Si hay más de un modelo, combinamos (merge) las columnas de Marca, Specs e Imagen
            if len(modelos) > 1:
                end_data_row = start_data_row + len(modelos) - 1
                
                ws.merge_cells(start_row=start_data_row, start_column=2, end_row=end_data_row, end_column=2) # Marca
                ws.merge_cells(start_row=start_data_row, start_column=4, end_row=end_data_row, end_column=4) # Specs
                ws.merge_cells(start_row=start_data_row, start_column=5, end_row=end_data_row, end_column=5) # Imagen
                
                # Re-asegurar alineación centrada para las celdas combinadas resultantes
                ws.cell(start_data_row, 2).alignment = Alignment(horizontal="center", vertical="center")
                ws.cell(start_data_row, 4).alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
                ws.cell(start_data_row, 5).alignment = Alignment(horizontal="center", vertical="center")
                
            current_row += len(modelos)
            
            # Dejar 3 filas en blanco
            current_row += 3
            
        # Ajustar anchos
        ws.column_dimensions["A"].width = 3
        ws.column_dimensions["B"].width = 25  # Marca
        ws.column_dimensions["C"].width = 20  # Modelos
        ws.column_dimensions["D"].width = 50  # Características técnicas
        ws.column_dimensions["E"].width = 30  # Imagen
        
        wb.save(str(out_path))
        _log("Planilla de fotos (datasheet) generada exitosamente.", "info", logger)
    except Exception as e:
        _log(f"Error al generar la planilla de fotos (datasheet): {e}", "error", logger)
        raise


# ═════════════════════════════════════════════════════════════════════════════
# 8. ORQUESTADOR
# ═════════════════════════════════════════════════════════════════════════════

def generate_solicitud(
    data: dict,
    oec: str,
    esquema: str = "",
    svg_bytes: bytes | None = None,
    logger: Any = None,
) -> dict:
    """
    Orquestador principal. Genera todos los archivos para Lenor o qetkra,
    los guarda en Solicitudes/[Nro]/ y retorna rutas + bytes del ZIP.
    """
    nro = data.get("certificado", "SIN_NRO").strip()
    if not nro:
        nro = "SIN_NRO"

    # Sanitizar número para nombres de archivos y directorios (reemplazar / \ : * ? " < > | por -)
    nro_clean = re.sub(r'[\\/:*?"<>|]', "-", nro)

    _log(f"--- ORQUESTADOR DE SOLICITUDES ---", "info", logger)
    _log(f"Parámetros recibidos: OEC={oec}, Certificado={nro} (sanitizado: {nro_clean}), Esquema={esquema}", "info", logger)

    # Normalizar OEC para qetkra
    oec_lower = oec.lower().strip()
    if oec_lower in ("quektra", "qetkra"):
        oec_lower = "qetkra"

    # Crear carpeta de salida
    out_dir = OUTPUT_BASE / nro_clean
    _log(f"Creando carpeta de salida (si no existe): {out_dir}", "debug", logger)
    out_dir.mkdir(parents=True, exist_ok=True)

    generated_files: list[Path] = []

    if oec_lower == "lenor":
        _log(f"Iniciando flujo de generación para organismo LENOR", "info", logger)
        xl_out = out_dir / f"Solicitud_Lenor_{nro_clean}.xlsm"
        generate_lenor_excel(data, xl_out, logger)
        generated_files.append(xl_out)

        docx_out = out_dir / f"Nota_Modelos_Lenor_{nro_clean}.docx"
        generate_lenor_word(data, docx_out, logger)
        generated_files.append(docx_out)

        ds_out = out_dir / f"Datasheet_{nro_clean}.xlsx"
        generate_lenor_datasheet_excel(data, ds_out, logger)
        generated_files.append(ds_out)

        if svg_bytes:
            _log("SVG de código QR provisto. Generando PDF con etiquetas QR vectoriales...", "info", logger)
            pdf_out = out_dir / f"QR Certificado {nro_clean}.pdf"
            generate_qr_pdf(svg_bytes, nro, pdf_out, logger)
            generated_files.append(pdf_out)
        else:
            _log("ADVERTENCIA: No se proveyó SVG de QR. Se omitirá la generación del PDF de etiquetas.", "warning", logger)

    elif oec_lower == "qetkra":
        _log(f"Iniciando flujo de generación para organismo QETKRA", "info", logger)
        xl_out = out_dir / f"Solicitud_qetkra_{nro_clean}.xlsx"
        generate_qetkra_excel(data, esquema, xl_out, logger)
        generated_files.append(xl_out)

        docx_out = out_dir / f"Nota_Correlatividad_qetkra_{nro_clean}.docx"
        generate_qetkra_word(data, docx_out, svg_bytes, logger)
        generated_files.append(docx_out)

    elif oec_lower in ("tuv", "tüv", "tuv rheinland"):
        _log(f"Iniciando flujo de generación para organismo TÜV RHEINLAND", "info", logger)
        docx_out = out_dir / f"Solicitud_tuv_{nro_clean}.docx"
        generate_tuv_word(data, docx_out, logger)
        generated_files.append(docx_out)

        ds_out = out_dir / f"Datasheet_{nro_clean}.xlsx"
        generate_lenor_datasheet_excel(data, ds_out, logger)
        generated_files.append(ds_out)

        if svg_bytes:
            _log("SVG de código QR provisto. Generando PDF con etiquetas QR vectoriales...", "info", logger)
            pdf_out = out_dir / f"QR Certificado {nro_clean}.pdf"
            generate_qr_pdf(svg_bytes, nro, pdf_out, logger)
            generated_files.append(pdf_out)
        else:
            _log("ADVERTENCIA: No se proveyó SVG de QR. Se omitirá la generación del PDF de etiquetas.", "warning", logger)

    else:
        _log(f"Error: Organismo '{oec}' desconocido.", "error", logger)
        raise ValueError(f"OEC desconocido: '{oec}'. Use 'lenor', 'qetkra' o 'tuv'.")

    # Comprimir en ZIP en memoria
    _log("Empaquetando archivos generados en formato ZIP...", "info", logger)
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in generated_files:
            _log(f"  + Agregando al ZIP: {fp.name}", "debug", logger)
            zf.write(fp, fp.name)
    zip_bytes = zip_buffer.getvalue()
    _log("Archivo ZIP creado correctamente en memoria.", "info", logger)
    _log(f"Generación para Certificado {nro} completada con éxito. Archivos: {[f.name for f in generated_files]}", "info", logger)

    return {
        "output_dir": str(out_dir),
        "files": [f.name for f in generated_files],
        "zip_bytes": zip_bytes,
    }
